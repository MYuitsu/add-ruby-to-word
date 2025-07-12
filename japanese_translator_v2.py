import google.generativeai as genai
import json
import re
from docx import Document
from docx.shared import RGBColor
import time
import os
from typing import List, Dict, Optional

# Regex patterns for Japanese text detection
JAPANESE_PATTERN = re.compile(r'[\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff]+')

class JapaneseToVietnameseTranslator:
    def __init__(self, api_key: str):
        """Initialize translator with Gemini API key"""
        if not api_key or api_key == "your_gemini_api_key_here":
            raise ValueError("Please provide a valid Gemini API key")
        
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')  # Sử dụng model mới hơn
        self.translation_cache = {}
        self.request_count = 0
        self.daily_request_count = 0  # Theo dõi requests hàng ngày
        self.max_requests_per_minute = 12  # Giới hạn API calls (free tier: 15/phút, để dư 3)
        self.max_requests_per_day = 45    # Giới hạn hàng ngày (free tier: 50/ngày, để dư 5)
        self.start_time = time.time()  # Theo dõi thời gian bắt đầu
        self.daily_start_time = time.time()  # Theo dõi thời gian bắt đầu ngày
        
    def has_japanese(self, text: str) -> bool:
        """Kiểm tra xem text có chứa ký tự tiếng Nhật không"""
        if not text:
            return False
        return bool(JAPANESE_PATTERN.search(text))
    
    def clean_japanese_text(self, text: str) -> str:
        """Làm sạch text tiếng Nhật"""
        # Loại bỏ whitespace thừa
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Loại bỏ các ký tự đặc biệt không cần thiết
        text = re.sub(r'[^\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff\s.,!?()（）「」『』、。！？]', '', text)
        
        return text
    
    def split_long_text(self, text: str, max_length: int = 500) -> List[str]:
        """Chia text dài thành các đoạn ngắn hơn để dịch"""
        if len(text) <= max_length:
            return [text]
        
        # Chia theo câu trước
        sentences = re.split(r'[。！？\n]', text)
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            if len(current_chunk + sentence) <= max_length:
                current_chunk += sentence + "。"
            else:
                if current_chunk:
                    chunks.append(current_chunk.rstrip("。"))
                current_chunk = sentence + "。"
        
        if current_chunk:
            chunks.append(current_chunk.rstrip("。"))
        
        return chunks
    
    def translate_text(self, japanese_text: str) -> Optional[str]:
        """Dịch text tiếng Nhật sang tiếng Việt bằng Gemini API"""
        if not japanese_text or not self.has_japanese(japanese_text):
            return None
        
        # Làm sạch text
        cleaned_text = self.clean_japanese_text(japanese_text)
        if not cleaned_text:
            return None
        
        # Kiểm tra cache trước
        cache_key = cleaned_text.strip()
        if cache_key in self.translation_cache:
            print(f"✓ Cache hit: {cache_key[:50]}...")
            return self.translation_cache[cache_key]
        
        # Kiểm tra rate limit với thời gian thực tế
        current_time = time.time()
        elapsed_time = current_time - self.start_time
        
        # Reset counter mỗi phút
        if elapsed_time >= 60:
            self.request_count = 0
            self.start_time = current_time
            elapsed_time = 0
        
        # Kiểm tra daily quota TRƯỚC
        if self.daily_request_count >= self.max_requests_per_day:
            print(f"🚫 DAILY QUOTA EXCEEDED! ({self.daily_request_count}/{self.max_requests_per_day})")
            print("   Free tier chỉ cho phép 50 requests/ngày")
            print("   Hãy thử lại vào ngày mai hoặc upgrade plan")
            return f"[Daily quota exceeded - {self.daily_request_count}/{self.max_requests_per_day}]"
        
        # Kiểm tra rate limit per minute
        if self.request_count >= self.max_requests_per_minute:
            wait_time = 60 - elapsed_time + 5  # Thêm 5 giây buffer
            print(f"⏳ Rate limit reached ({self.request_count}/{self.max_requests_per_minute}), waiting {wait_time:.1f} seconds...")
            time.sleep(wait_time)
            self.request_count = 0
            self.start_time = time.time()
        
        try:
            # Chia text dài thành chunks nhỏ hơn
            chunks = self.split_long_text(cleaned_text)
            translations = []
            
            for chunk in chunks:
                prompt = f"""
Dịch đoạn văn tiếng Nhật sau sang tiếng Việt một cách tự nhiên và chính xác:

Yêu cầu:
- Dịch chính xác nghĩa và ngữ cảnh
- Sử dụng từ ngữ tiếng Việt tự nhiên, không máy móc
- Giữ nguyên cấu trúc câu hợp lý
- Không thêm giải thích hay chú thích
- Chỉ trả về bản dịch tiếng Việt

Văn bản tiếng Nhật:
{chunk}

Bản dịch tiếng Việt:"""

                response = self.model.generate_content(
                    prompt,
                    generation_config={
                        'temperature': 0.3,  # Giảm tính ngẫu nhiên
                        'max_output_tokens': 2048,
                    }
                )
                
                if response.text:
                    translation = response.text.strip()
                    # Loại bỏ các prefix không mong muốn
                    translation = re.sub(r'^(Bản dịch tiếng Việt:|Dịch:|Translation:)\s*', '', translation, flags=re.IGNORECASE)
                    translations.append(translation)
                
                self.request_count += 1
                self.daily_request_count += 1  # Cập nhật daily counter
                
                # Delay giữa các request để tránh rate limit (tăng delay)
                time.sleep(1.0)  # Tăng từ 0.5 lên 1.0 giây
            
            # Kết hợp các translations
            final_translation = " ".join(translations)
            
            # Lưu vào cache
            self.translation_cache[cache_key] = final_translation
            
            print(f"✓ Translated: {cleaned_text[:50]}... → {final_translation[:50]}...")
            return final_translation
            
        except Exception as e:
            error_msg = str(e)
            
            # Kiểm tra nếu là daily quota error
            if "429" in error_msg and "FreeTier" in error_msg and "PerDay" in error_msg:
                print(f"🚫 DAILY QUOTA EXCEEDED! API đã vượt quá 50 requests/ngày")
                print("   Giải pháp:")
                print("   1. Đợi đến ngày mai để reset quota")
                print("   2. Upgrade lên paid plan tại: https://aistudio.google.com/app/pricing")
                print("   3. Sử dụng cache cho các text đã dịch")
                return f"[Daily quota exceeded - Hãy thử lại vào ngày mai]"
            
            # Các lỗi khác
            error_msg = f"[Lỗi dịch: {str(e)}]"
            print(f"✗ Error translating '{cleaned_text[:50]}...': {e}")
            return error_msg
    
    def add_translation_to_paragraph(self, paragraph, vietnamese_text: str):
        """Thêm bản dịch tiếng Việt vào paragraph"""
        # Thêm xuống dòng
        paragraph.add_run().add_break()
        
        # Thêm text tiếng Việt với format đặc biệt
        vn_run = paragraph.add_run(f"🇻🇳 {vietnamese_text}")
        
        # Định dạng cho text tiếng Việt
        vn_run.font.italic = True
        vn_run.font.color.rgb = RGBColor(0, 100, 0)  # Màu xanh lá đậm
        vn_run.font.size = vn_run.font.size  # Giữ nguyên size
    
    def process_paragraph(self, paragraph) -> bool:
        """Xử lý một paragraph và thêm bản dịch tiếng Việt"""
        original_text = paragraph.text.strip()
        
        if not original_text or not self.has_japanese(original_text):
            return False
        
        # Dịch text
        vietnamese_translation = self.translate_text(original_text)
        
        if not vietnamese_translation or vietnamese_translation.startswith("[Lỗi dịch"):
            return False
        
        # Thêm bản dịch vào paragraph
        self.add_translation_to_paragraph(paragraph, vietnamese_translation)
        
        return True
    
    def process_table_cell(self, cell) -> int:
        """Xử lý các paragraph trong table cell"""
        translated_count = 0
        
        for paragraph in cell.paragraphs:
            if self.process_paragraph(paragraph):
                translated_count += 1
        
        return translated_count
    
    def process_word_document(self, input_path: str, output_path: str):
        """Xử lý file Word và thêm bản dịch tiếng Việt"""
        print(f"🔄 Đang xử lý tài liệu: {input_path}")
        start_time = time.time()
        
        try:
            doc = Document(input_path)
            
            # Thống kê
            stats = {
                'total_paragraphs': 0,
                'japanese_paragraphs': 0,
                'translated_paragraphs': 0,
                'tables_processed': 0,
                'cells_processed': 0
            }
            
            print("📄 Đang xử lý paragraphs chính...")
            
            # Xử lý paragraphs chính
            for i, paragraph in enumerate(doc.paragraphs):
                stats['total_paragraphs'] += 1
                
                if self.has_japanese(paragraph.text):
                    stats['japanese_paragraphs'] += 1
                    
                    if self.process_paragraph(paragraph):
                        stats['translated_paragraphs'] += 1
                    
                    # Hiển thị tiến trình
                    if stats['japanese_paragraphs'] % 10 == 0:
                        elapsed = time.time() - start_time
                        print(f"  📝 Đã xử lý {stats['japanese_paragraphs']} đoạn tiếng Nhật - {elapsed:.1f}s")
            
            print("📊 Đang xử lý tables...")
            
            # Xử lý tables
            for table_idx, table in enumerate(doc.tables):
                stats['tables_processed'] += 1
                
                for row in table.rows:
                    for cell in row.cells:
                        stats['cells_processed'] += 1
                        cell_translated = self.process_table_cell(cell)
                        stats['translated_paragraphs'] += cell_translated
                
                # Hiển thị tiến trình table
                if (table_idx + 1) % 5 == 0:
                    elapsed = time.time() - start_time
                    print(f"  📋 Đã xử lý {table_idx + 1} tables - {elapsed:.1f}s")
            
            # Lưu tài liệu
            print("💾 Đang lưu tài liệu...")
            save_start = time.time()
            doc.save(output_path)
            save_time = time.time() - save_start
            
            # Lưu cache
            cache_file = output_path.replace('.docx', '_translation_cache.json')
            self.save_cache(cache_file)
            
            total_time = time.time() - start_time
            
            # In báo cáo kết quả
            self.print_results(output_path, stats, total_time, save_time, cache_file)
            
        except Exception as e:
            print(f"❌ Lỗi xử lý tài liệu: {e}")
            import traceback
            traceback.print_exc()
    
    def save_cache(self, cache_file: str):
        """Lưu cache dịch thuật"""
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.translation_cache, f, ensure_ascii=False, indent=2)
            print(f"💾 Đã lưu {len(self.translation_cache)} bản dịch vào cache")
        except Exception as e:
            print(f"❌ Lỗi lưu cache: {e}")
    
    def load_cache(self, cache_file: str):
        """Tải cache dịch thuật có sẵn"""
        try:
            if os.path.exists(cache_file):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    self.translation_cache = json.load(f)
                print(f"📥 Đã tải {len(self.translation_cache)} bản dịch từ cache")
            else:
                print("📝 Không tìm thấy cache, bắt đầu từ đầu")
        except Exception as e:
            print(f"❌ Lỗi tải cache: {e}")
    
    def print_results(self, output_path: str, stats: Dict, total_time: float, save_time: float, cache_file: str):
        """In báo cáo kết quả"""
        print("\n" + "="*60)
        print("🎉 KẾT QUẢ DỊCH THUẬT")
        print("="*60)
        print(f"📁 File đầu ra: {output_path}")
        print(f"📊 Tổng paragraphs: {stats['total_paragraphs']}")
        print(f"🇯🇵 Paragraphs tiếng Nhật: {stats['japanese_paragraphs']}")
        print(f"🇻🇳 Paragraphs đã dịch: {stats['translated_paragraphs']}")
        print(f"📋 Tables đã xử lý: {stats['tables_processed']}")
        print(f"📄 Cells đã xử lý: {stats['cells_processed']}")
        print(f"💾 Bản dịch trong cache: {len(self.translation_cache)}")
        print(f"⏱️  Thời gian xử lý: {total_time:.2f} giây")
        print(f"💾 Thời gian lưu: {save_time:.2f} giây")
        print(f"📂 Cache file: {cache_file}")
        
        if stats['japanese_paragraphs'] > 0:
            success_rate = (stats['translated_paragraphs'] / stats['japanese_paragraphs']) * 100
            print(f"✅ Tỷ lệ thành công: {success_rate:.1f}%")
        
        print("="*60)
    
    def batch_translate_paragraphs(self, paragraphs_data: List[tuple], batch_size: int = 5) -> Dict[str, str]:
        """Dịch nhiều đoạn văn cùng lúc để tiết kiệm API calls"""
        batch_results = {}
        
        # Tạo batches từ paragraphs chưa có trong cache
        batches = []
        current_batch = []
        
        for para_id, text in paragraphs_data:
            cleaned_text = self.clean_japanese_text(text)
            if cleaned_text and cleaned_text not in self.translation_cache:
                current_batch.append((para_id, cleaned_text))
                
                if len(current_batch) >= batch_size:
                    batches.append(current_batch)
                    current_batch = []
        
        if current_batch:
            batches.append(current_batch)
        
        # Xử lý từng batch
        for batch_idx, batch in enumerate(batches):
            print(f"🔄 Processing batch {batch_idx + 1}/{len(batches)} ({len(batch)} items)")
            
            # Tạo combined prompt cho cả batch
            combined_text = "\n---\n".join([f"[{para_id}] {text}" for para_id, text in batch])
            
            prompt = f"""
Dịch các đoạn văn tiếng Nhật sau sang tiếng Việt. Mỗi đoạn được đánh số [ID].
Trả về kết quả theo định dạng: [ID] Bản dịch tiếng Việt

Yêu cầu:
- Dịch chính xác nghĩa và ngữ cảnh
- Sử dụng từ ngữ tiếng Việt tự nhiên
- Giữ đúng định dạng [ID] trước mỗi bản dịch

Văn bản tiếng Nhật:
{combined_text}

Bản dịch tiếng Việt:"""

            try:
                # Kiểm tra rate limit trước khi gọi API
                current_time = time.time()
                elapsed_time = current_time - self.start_time
                
                if elapsed_time >= 60:
                    self.request_count = 0
                    self.start_time = current_time
                elif self.request_count >= self.max_requests_per_minute:
                    wait_time = 60 - elapsed_time + 5
                    print(f"⏳ Waiting {wait_time:.1f}s for rate limit...")
                    time.sleep(wait_time)
                    self.request_count = 0
                    self.start_time = time.time()
                
                response = self.model.generate_content(
                    prompt,
                    generation_config={
                        'temperature': 0.3,
                        'max_output_tokens': 2048,
                    }
                )
                
                self.request_count += 1
                
                if response.text:
                    # Parse kết quả batch
                    lines = response.text.strip().split('\n')
                    for line in lines:
                        line = line.strip()
                        if line.startswith('[') and ']' in line:
                            try:
                                end_bracket = line.index(']')
                                para_id = line[1:end_bracket]
                                translation = line[end_bracket + 1:].strip()
                                
                                # Tìm text gốc tương ứng
                                for orig_id, orig_text in batch:
                                    if orig_id == para_id:
                                        self.translation_cache[orig_text] = translation
                                        batch_results[orig_text] = translation
                                        break
                            except (ValueError, IndexError):
                                continue
                
                # Delay giữa các batch
                time.sleep(2.0)
                
            except Exception as e:
                print(f"❌ Error in batch {batch_idx + 1}: {e}")
                # Fallback: dịch từng đoạn riêng
                for para_id, text in batch:
                    try:
                        translation = self.translate_text(text)
                        if translation:
                            batch_results[text] = translation
                    except Exception as fallback_error:
                        print(f"❌ Fallback error for {para_id}: {fallback_error}")
        
        return batch_results

def main():
    """Hàm chính"""
    # Import config từ file riêng
    try:
        from config import get_config, print_config
        CONFIG = get_config()
        if CONFIG is None:
            return
        print_config()
    except ImportError:
        print("❌ Không tìm thấy file config.py!")
        print("   Vui lòng tạo file config.py với API key")
        return
    
    print("🌟 JAPANESE TO VIETNAMESE TRANSLATOR V2")
    print("="*50)
    print(f"📥 Input: {CONFIG['input_file']}")
    print(f"📤 Output: {CONFIG['output_file']}")
    print()
    
    # Kiểm tra file input
    if not os.path.exists(CONFIG['input_file']):
        print(f"❌ Không tìm thấy file input: {CONFIG['input_file']}")
        return
    
    try:
        # Khởi tạo translator
        print("🔧 Đang khởi tạo translator...")
        translator = JapaneseToVietnameseTranslator(CONFIG['api_key'])
        
        # Tải cache có sẵn
        translator.load_cache(CONFIG['cache_file'])
        
        # Xử lý tài liệu
        translator.process_word_document(CONFIG['input_file'], CONFIG['output_file'])
        
        print("\n🎊 DỊCH THUẬT HOÀN THÀNH!")
        
    except Exception as e:
        print(f"❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
