#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Script để xóa tất cả tiếng Việt khỏi file Word, chỉ giữ lại tiếng Nhật và tiếng Anh
"""

import re
import unicodedata
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import os
import sys

def is_vietnamese_char(char):
    """Kiểm tra xem ký tự có phải là ký tự tiếng Việt không"""
    vietnamese_chars = set([
        'à', 'á', 'ả', 'ã', 'ạ', 'ă', 'ằ', 'ắ', 'ẳ', 'ẵ', 'ặ',
        'â', 'ầ', 'ấ', 'ẩ', 'ẫ', 'ậ', 'è', 'é', 'ẻ', 'ẽ', 'ẹ',
        'ê', 'ề', 'ế', 'ể', 'ễ', 'ệ', 'ì', 'í', 'ỉ', 'ĩ', 'ị',
        'ò', 'ó', 'ỏ', 'õ', 'ọ', 'ô', 'ồ', 'ố', 'ổ', 'ỗ', 'ộ',
        'ơ', 'ờ', 'ớ', 'ở', 'ỡ', 'ợ', 'ù', 'ú', 'ủ', 'ũ', 'ụ',
        'ư', 'ừ', 'ứ', 'ử', 'ữ', 'ự', 'ỳ', 'ý', 'ỷ', 'ỹ', 'ỵ',
        'đ', 'À', 'Á', 'Ả', 'Ã', 'Ạ', 'Ă', 'Ằ', 'Ắ', 'Ẳ', 'Ẵ', 'Ặ',
        'Â', 'Ầ', 'Ấ', 'Ẩ', 'Ẫ', 'Ậ', 'È', 'É', 'Ẻ', 'Ẽ', 'Ẹ',
        'Ê', 'Ề', 'Ế', 'Ể', 'Ễ', 'Ệ', 'Ì', 'Í', 'Ỉ', 'Ĩ', 'Ị',
        'Ò', 'Ó', 'Ỏ', 'Õ', 'Ọ', 'Ô', 'Ồ', 'Ố', 'Ổ', 'Ỗ', 'Ộ',
        'Ơ', 'Ờ', 'Ớ', 'Ở', 'Ỡ', 'Ợ', 'Ù', 'Ú', 'Ủ', 'Ũ', 'Ụ',
        'Ư', 'Ừ', 'Ứ', 'Ử', 'Ữ', 'Ự', 'Ỳ', 'Ý', 'Ỷ', 'Ỹ', 'Ỵ',
        'Đ'
    ])
    return char in vietnamese_chars

def is_japanese_char(char):
    """Kiểm tra xem ký tự có phải là tiếng Nhật không"""
    # Hiragana: U+3040-U+309F
    # Katakana: U+30A0-U+30FF
    # CJK Unified Ideographs: U+4E00-U+9FFF (Kanji)
    # CJK Symbols and Punctuation: U+3000-U+303F
    code = ord(char)
    return (0x3040 <= code <= 0x309F or  # Hiragana
            0x30A0 <= code <= 0x30FF or  # Katakana
            0x4E00 <= code <= 0x9FFF or  # Kanji
            0x3000 <= code <= 0x303F)    # CJK punctuation

def is_english_char(char):
    """Kiểm tra xem ký tự có phải là tiếng Anh không"""
    return char.isascii() and (char.isalpha() or char.isdigit() or char in ' .,!?;:()[]{}"-\'')

# Biến global để lưu danh sách từ tiếng Việt
_vietnamese_words_set = None

def load_vietnamese_words():
    """Tải danh sách từ tiếng Việt từ file vi-DauMoi.txt"""
    global _vietnamese_words_set
    
    if _vietnamese_words_set is not None:
        return _vietnamese_words_set
    
    vietnamese_words = set()
    
    try:
        with open('vi-DauMoi.txt', 'r', encoding='utf-8') as f:
            for line in f:
                word = line.strip()
                if word and len(word) > 0:
                    # Thêm từ gốc
                    vietnamese_words.add(word.lower())
                    
                    # Nếu từ có dấu, thêm cả phiên bản không dấu
                    word_no_accent = remove_accents(word)
                    if word_no_accent != word:
                        vietnamese_words.add(word_no_accent.lower())
        
        print(f"✅ Đã tải {len(vietnamese_words)} từ tiếng Việt từ file vi-DauMoi.txt")
        
    except FileNotFoundError:
        print("⚠️  Không tìm thấy file vi-DauMoi.txt, sử dụng từ điển hardcode")
        # Fallback về từ điển hardcode nếu không tìm thấy file
        vietnamese_words = get_hardcoded_vietnamese_words()
    except Exception as e:
        print(f"⚠️  Lỗi khi đọc file vi-DauMoi.txt: {e}, sử dụng từ điển hardcode")
        vietnamese_words = get_hardcoded_vietnamese_words()
    
    _vietnamese_words_set = vietnamese_words
    return vietnamese_words

def remove_accents(text):
    """Loại bỏ dấu từ text tiếng Việt"""
    # Bảng chuyển đổi dấu tiếng Việt
    vietnamese_accent_map = {
        'à': 'a', 'á': 'a', 'ả': 'a', 'ã': 'a', 'ạ': 'a',
        'ă': 'a', 'ằ': 'a', 'ắ': 'a', 'ẳ': 'a', 'ẵ': 'a', 'ặ': 'a',
        'â': 'a', 'ầ': 'a', 'ấ': 'a', 'ẩ': 'a', 'ẫ': 'a', 'ậ': 'a',
        'è': 'e', 'é': 'e', 'ẻ': 'e', 'ẽ': 'e', 'ẹ': 'e',
        'ê': 'e', 'ề': 'e', 'ế': 'e', 'ể': 'e', 'ễ': 'e', 'ệ': 'e',
        'ì': 'i', 'í': 'i', 'ỉ': 'i', 'ĩ': 'i', 'ị': 'i',
        'ò': 'o', 'ó': 'o', 'ỏ': 'o', 'õ': 'o', 'ọ': 'o',
        'ô': 'o', 'ồ': 'o', 'ố': 'o', 'ổ': 'o', 'ỗ': 'o', 'ộ': 'o',
        'ơ': 'o', 'ờ': 'o', 'ớ': 'o', 'ở': 'o', 'ỡ': 'o', 'ợ': 'o',
        'ù': 'u', 'ú': 'u', 'ủ': 'u', 'ũ': 'u', 'ụ': 'u',
        'ư': 'u', 'ừ': 'u', 'ứ': 'u', 'ử': 'u', 'ữ': 'u', 'ự': 'u',
        'ỳ': 'y', 'ý': 'y', 'ỷ': 'y', 'ỹ': 'y', 'ỵ': 'y',
        'đ': 'd',
        'À': 'A', 'Á': 'A', 'Ả': 'A', 'Ã': 'A', 'Ạ': 'A',
        'Ă': 'A', 'Ằ': 'A', 'Ắ': 'A', 'Ẳ': 'A', 'Ẵ': 'A', 'Ặ': 'A',
        'Â': 'A', 'Ầ': 'A', 'Ấ': 'A', 'Ẩ': 'A', 'Ẫ': 'A', 'Ậ': 'A',
        'È': 'E', 'É': 'E', 'Ẻ': 'E', 'Ẽ': 'E', 'Ẹ': 'E',
        'Ê': 'E', 'Ề': 'E', 'Ế': 'E', 'Ể': 'E', 'Ễ': 'E', 'Ệ': 'E',
        'Ì': 'I', 'Í': 'I', 'Ỉ': 'I', 'Ĩ': 'I', 'Ị': 'I',
        'Ò': 'O', 'Ó': 'O', 'Ỏ': 'O', 'Õ': 'O', 'Ọ': 'O',
        'Ô': 'O', 'Ồ': 'O', 'Ố': 'O', 'Ổ': 'O', 'Ỗ': 'O', 'Ộ': 'O',
        'Ơ': 'O', 'Ờ': 'O', 'Ớ': 'O', 'Ở': 'O', 'Ỡ': 'O', 'Ợ': 'O',
        'Ù': 'U', 'Ú': 'U', 'Ủ': 'U', 'Ũ': 'U', 'Ụ': 'U',
        'Ư': 'U', 'Ừ': 'U', 'Ứ': 'U', 'Ử': 'U', 'Ữ': 'U', 'Ự': 'U',
        'Ỳ': 'Y', 'Ý': 'Y', 'Ỷ': 'Y', 'Ỹ': 'Y', 'Ỵ': 'Y',
        'Đ': 'D'
    }
    
    result = ''
    for char in text:
        if char in vietnamese_accent_map:
            result += vietnamese_accent_map[char]
        else:
            result += char
    return result

def get_hardcoded_vietnamese_words():
    """Trả về từ điển hardcode làm fallback"""
    return {
        # Từ điển hardcode cơ bản
        'toi', 'ban', 'no', 'chung', 'ta', 'minh', 'ho', 'may', 'cac', 'nhung', 'moi',
        'nguoi', 'ai', 'gi', 'dau', 'nao', 'sao', 'bao', 'tat', 'ca', 'mang',
        'la', 'co', 'di', 'den', 'lam', 'noi', 'biet', 'thay', 'nghe', 'doc', 'viet', 'hoc',
        'an', 'uong', 'ngu', 'choi', 'xem', 'mua', 'ban', 'cho', 'nhan', 'gui', 'goi',
        'sinh', 'theo', 'van', 'con', 'duoc', 'tim', 'hieu', 'dong', 'chia', 'nhan',
        'dung', 'ngoi', 'nam', 'chay', 'nhay', 'hat', 'khoc', 'cuoi', 'ket',
        'bat', 'dau', 'thuc', 'mo', 'dong', 'cat', 'gan', 'roi', 'bo', 'lay',
        'dem', 'mang', 'dua', 'dieu', 'khien', 'quan', 'ly', 'tri', 'chua',
        'benh', 'viec', 'hoi', 'dap', 'thi', 'tot', 'dep', 'xau', 'cao', 'thap',
        'dai', 'ngan', 'rong', 'hep', 'lon', 'nho', 'nong', 'lanh', 'nhanh', 'cham',
        'khoe', 'om', 'vui', 'buon', 'gioi', 'te', 'trung', 'binh', 'sai', 'dung',
        'chan', 'le', 'cu', 'tre', 'gia', 'khac', 'giong', 'bang', 'hon', 'kem',
        'nhat', 'nhi', 'ba', 'cuoi', 'xa', 'mat', 'am', 'kho', 'uot', 'de',
        'khan', 'den', 'trang', 'do', 'xanh', 'vang', 'tim', 'hong', 'nau',
        'nha', 'truong', 'lop', 'ghe', 'sach', 'but', 'giay', 'ao', 'quan',
        'me', 'bo', 'anh', 'chi', 'em', 'ong', 'ba', 'thay', 'chu',
        'nuoc', 'com', 'pho', 'banh', 'ca', 'thit', 'rau', 'trai', 'gao', 'mi',
        'xe', 'may', 'tau', 'oto', 'duong', 'hang', 'tien', 'tra', 'bieu', 'mau',
        'bai', 'tap', 'de', 'giua', 'phai', 'ben', 'canh', 'tren', 'duoi',
        'trong', 'ngoai', 'truoc', 'sau', 'day', 'kia', 'nay', 'mai', 'hom', 'qua',
        'tuoi', 'nu', 'dan', 'em', 'mot', 'hai', 'bon', 'sau', 'bay', 'tam',
        'chin', 'muoi', 'tram', 'ngan', 'trieu', 'ty', 'le', 'lan', 'phan',
        'ngay', 'thang', 'gio', 'phut', 'giay', 'tuan', 'thoi', 'gian', 'som',
        'muon', 'dem', 'sang', 'chieu', 'toi', 'khuya', 'trua', 'tet', 'hoi',
        'va', 'hoac', 'ma', 'neu', 'vi', 'de', 'tu', 'ben', 'boi', 'nhu',
        'nham', 'vay', 'the', 'khi', 'luc', 'gio', 'tai', 'o', 'ra', 'cach',
        'rat', 'qua', 'cung', 'da', 'dang', 'se', 'chua', 'roi', 'khong',
        'cai', 'chiec', 'chuyen', 'dieu', 'su', 'ay', 'do', 'het', 'xong',
        'co', 'the', 'nen', 'phai', 'can', 'pha', 'i', 'muon', 'thich', 'ghet',
        'so', 'yeu', 'thuong', 'nho', 'quen'
    }

def is_vietnamese_word(word):
    """Kiểm tra xem từ có phải là tiếng Việt không (sử dụng từ điển từ file)"""
    vietnamese_words = load_vietnamese_words()
    word_lower = word.lower().strip('.,!?;:()[]{}"-\'')
    return word_lower in vietnamese_words

def is_vietnamese_text(text):
    """Kiểm tra xem đoạn text có phải là tiếng Việt không"""
    if not text or not text.strip():
        return False
    
    # Đếm số ký tự tiếng Việt (có dấu)
    vietnamese_char_count = sum(1 for char in text if is_vietnamese_char(char))
    
    # Đếm số từ tiếng Việt (không dấu)
    words = text.split()
    vietnamese_word_count = sum(1 for word in words if is_vietnamese_word(word))
    
    total_chars = len([char for char in text if char.isalpha()])
    total_words = len([word for word in words if word.strip('.,!?;:()[]{}"-\'').isalpha()])
    
    if total_chars == 0 and total_words == 0:
        return False
    
    # Nếu có ký tự tiếng Việt có dấu
    if total_chars > 0 and vietnamese_char_count / total_chars > 0.1:
        return True
    
    # Nếu có từ tiếng Việt không dấu
    if total_words > 0 and vietnamese_word_count / total_words > 0.3:
        return True
    
    return False

def should_remove_vietnamese_word(word):
    """Kiểm tra xem từ có nên bị xóa (là từ tiếng Việt) không"""
    # Loại bỏ dấu câu để kiểm tra từ thuần
    clean_word = word.strip('.,!?;:()[]{}"-\'')
    
    # Nếu từ rỗng sau khi loại bỏ dấu câu
    if not clean_word:
        return False
    
    # Nếu từ có ký tự tiếng Việt có dấu thì xóa
    if any(is_vietnamese_char(char) for char in clean_word):
        return True
    
    # Nếu từ là từ tiếng Việt không dấu thì xóa
    if is_vietnamese_word(clean_word):
        return True
    
    # Kiểm tra xem từ có phải là từ tiếng Việt viết tắt hoặc ghép không
    # Ví dụ: "cht", "lng", "qun", "phng", v.v.
    if len(clean_word) >= 2 and clean_word.lower() in load_vietnamese_words():
        return True
    
    # Danh sách các từ viết tắt tiếng Việt thường gặp
    vietnamese_abbreviations = {
        'mm', 'cht', 'lng', 'qun', 'phng', 'thn', 'mc', 'sn', 'ph', 'bt', 'ct', 'tr', 'ng', 
        'nh', 'th', 'kh', 'ch', 'vt', 'mt', 'ht', 'dt', 'gm', 'tm', 'nm', 'pt', 'st',
        'lt', 'gt', 'dn', 'cn', 'tn', 'bn', 'pn', 'mn', 'hn', 'gn', 'fn', 'rn',
        'cch', 'kch', 'thc', 'phc', 'nht', 'tht', 'sht', 'ght', 'dht', 'bht',
        'cc', 'dc', 'tc', 'nc', 'pc', 'gc', 'fc', 'bc', 'hc', 'lc', 'rc', 'sc',
        'tng', 'dng', 'bng', 'png', 'mng', 'lng', 'rng', 'sng', 'hng', 'gng',
        'cp', 'tp', 'dp', 'bp', 'hp', 'lp', 'rp', 'sp', 'mp', 'np', 'gp', 'fp',
        'cm', 'tm', 'dm', 'bm', 'hm', 'lm', 'rm', 'sm', 'nm', 'gm', 'fm', 'pm',
        # Các từ viết tắt bổ sung
        'mmt', 'mmm', 'trt', 'sst', 'bbt', 'lll', 'kkk', 'rrr', 'nnn', 'ttt',
        'l', 'k', 'n', 'r', 's', 'b', 'g', 'f', 'h', 'j', 'p', 'q', 'v', 'w', 'x', 'z',
        # Từ viết tắt của "khái niệm", "kiểm soát", "chất lượng"
        'kni', 'ksot', 'clung', 'khai', 'niem', 'kiem', 'soat', 'chat', 'luong',
        'cqun', 'ccht', 'clng', 'qll', 'qlt', 'clt', 'qnl', 'knl', 'ksl'
    }
    
    # Kiểm tra từ viết tắt tiếng Việt
    if clean_word.lower() in vietnamese_abbreviations:
        return True
    
    # Kiểm tra các ký tự đơn lẻ có thể là viết tắt tiếng Việt
    if len(clean_word) == 1:
        # Những ký tự đơn lẻ thường là viết tắt tiếng Việt (trừ A, I)
        single_chars_to_remove = set('bcdfghjklmnpqrstvwxyz')
        if clean_word.lower() in single_chars_to_remove:
            return True
    
    # Kiểm tra các pattern tiếng Việt không dấu
    # Pattern 1: từ chỉ có phụ âm (không có nguyên âm a,e,i,o,u)
    if len(clean_word) >= 2 and len(clean_word) <= 4:
        vowels = set('aeiouAEIOU')
        if not any(char in vowels for char in clean_word):
            # Kiểm tra xem có phải là từ tiếng Anh hay không
            common_english_consonant_words = {'by', 'my', 'try', 'fly', 'sky', 'dry', 'cry', 'fry', 'why'}
            if clean_word.lower() not in common_english_consonant_words:
                return True
    
    return False

def should_keep_paragraph(text):
    """Quyết định có nên giữ lại paragraph không"""
    if not text or not text.strip():
        return False
    
    # Nếu có ký tự Nhật thì giữ lại
    if any(is_japanese_char(char) for char in text):
        return True
    
    # Đếm số từ tiếng Việt trong đoạn text
    words = text.split()
    vietnamese_word_count = 0
    total_words = 0
    
    for word in words:
        # Bỏ qua dấu câu khi đếm từ
        clean_word = word.strip('.,!?;:()[]{}"-\'')
        if clean_word and clean_word.isalpha():
            total_words += 1
            if should_remove_vietnamese_word(word):
                vietnamese_word_count += 1
    
    # Nếu không có từ nào thì giữ lại
    if total_words == 0:
        return True
    
    # Nếu có hơn 30% từ tiếng Việt thì xóa cả đoạn
    vietnamese_ratio = vietnamese_word_count / total_words
    if vietnamese_ratio > 0.3:
        return False
    
    # Nếu chủ yếu là tiếng Anh thì giữ lại
    alpha_chars = [char for char in text if char.isalpha()]
    if alpha_chars:
        english_count = sum(1 for char in alpha_chars if is_english_char(char))
        if english_count / len(alpha_chars) > 0.6:
            return True
    
    # Mặc định giữ lại nếu không có nhiều từ tiếng Việt
    return True

def clean_text(text):
    """Làm sạch text, xóa các từ tiếng Việt"""
    if not text:
        return ""
    
    # Tách thành các từ
    words = text.split()
    clean_words = []
    
    for word in words:
        # Nếu từ có ký tự Nhật thì giữ nguyên
        if any(is_japanese_char(char) for char in word):
            clean_words.append(word)
        # Nếu từ là tiếng Việt (có dấu hoặc không dấu) thì bỏ qua hoàn toàn
        elif should_remove_vietnamese_word(word):
            continue  # Bỏ qua từ này hoàn toàn
        # Nếu từ chủ yếu là tiếng Anh hoặc ký tự khác thì giữ nguyên
        else:
            clean_words.append(word)
    
    return ' '.join(clean_words)

def clean_text_preserve_format(text):
    """Làm sạch text nhưng giữ nguyên vị trí để bảo tồn format"""
    if not text:
        return ""
    
    # Tách thành các từ nhưng giữ nguyên khoảng trắng
    words = text.split(' ')
    clean_words = []
    
    for word in words:
        if not word:  # Khoảng trắng trống
            clean_words.append(word)
            continue
            
        # Nếu từ có ký tự Nhật thì giữ nguyên
        if any(is_japanese_char(char) for char in word):
            clean_words.append(word)
        # Nếu từ là tiếng Việt thì BỎ QUA HOÀN TOÀN (không thêm gì cả)
        elif should_remove_vietnamese_word(word):
            continue  # Bỏ qua từ này hoàn toàn, không thêm vào clean_words
        # Nếu từ chủ yếu là tiếng Anh hoặc ký tự khác thì giữ nguyên
        else:
            clean_words.append(word)
    
    # Nối lại và loại bỏ khoảng trắng thừa
    result = ' '.join(clean_words)
    # Loại bỏ nhiều khoảng trắng liên tiếp
    result = re.sub(r'\s+', ' ', result)
    return result.strip()

def clean_run_preserve_format(run):
    """Làm sạch một run trong Word document và giữ nguyên format"""
    if not run.text:
        return False
    
    original_text = run.text
    cleaned_text = clean_text_preserve_format(original_text)
    
    # Nếu text bị thay đổi thì cập nhật
    if cleaned_text != original_text:
        run.text = cleaned_text
        return True
    
    return False

def remove_vietnamese_from_docx(input_file, output_file):
    """Xóa tiếng Việt khỏi file Word và giữ nguyên format"""
    print(f"📖 Đang xử lý file: {input_file}")
    
    try:
        # Mở file Word
        doc = Document(input_file)
        
        # Thống kê
        total_paragraphs = len(doc.paragraphs)
        removed_paragraphs = 0
        processed_paragraphs = 0
        cleaned_paragraphs = 0
        cleaned_runs = 0
        
        print(f"📊 Tổng số paragraph: {total_paragraphs}")
        
        # Xử lý từng paragraph
        paragraphs_to_remove = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            original_text = paragraph.text
            
            if not original_text.strip():
                continue
                
            processed_paragraphs += 1
            
            # Kiểm tra xem có nên giữ lại paragraph không
            if should_keep_paragraph(original_text):
                # Làm sạch toàn bộ text của paragraph trước
                cleaned_full_text = clean_text_preserve_format(original_text)
                
                # Nếu text đã thay đổi, cập nhật lại toàn bộ paragraph
                if cleaned_full_text != original_text:
                    # Xóa text khỏi tất cả runs hiện tại và chỉ giữ lại run đầu tiên
                    if paragraph.runs:
                        # Giữ run đầu tiên và xóa text của các run khác
                        first_run = paragraph.runs[0]
                        first_run.text = cleaned_full_text
                        
                        # Xóa text của các run còn lại
                        for run in paragraph.runs[1:]:
                            run.text = ""
                    else:
                        # Nếu không có run nào, tạo run mới
                        new_run = paragraph.add_run()
                        new_run.text = cleaned_full_text
                    
                    cleaned_paragraphs += 1
                    cleaned_runs += 1
                    print(f"✅ Paragraph {i+1}: Đã làm sạch toàn bộ")
                    print(f"   Trước: {original_text[:100]}...")
                    print(f"   Sau:  {cleaned_full_text[:100]}...")
                
                # Kiểm tra xem paragraph còn nội dung không sau khi làm sạch
                if not cleaned_full_text.strip():
                    paragraphs_to_remove.append(paragraph)
                    removed_paragraphs += 1
                    print(f"❌ Paragraph {i+1}: Xóa (không còn nội dung sau khi làm sạch)")
            else:
                # Đánh dấu để xóa toàn bộ paragraph
                paragraphs_to_remove.append(paragraph)
                removed_paragraphs += 1
                print(f"❌ Paragraph {i+1}: Xóa (có nhiều từ tiếng Việt)")
                print(f"   Nội dung: {original_text[:100]}...")
        
        # Xóa các paragraph đã đánh dấu
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Lưu file kết quả
        doc.save(output_file)
        
        print(f"\n🎉 Hoàn thành!")
        print(f"📊 Thống kê:")
        print(f"   - Tổng số paragraph: {total_paragraphs}")
        print(f"   - Đã xử lý: {processed_paragraphs}")
        print(f"   - Đã làm sạch: {cleaned_paragraphs}")
        print(f"   - Đã làm sạch runs: {cleaned_runs}")
        print(f"   - Đã xóa: {removed_paragraphs}")
        print(f"   - Còn lại: {total_paragraphs - removed_paragraphs}")
        print(f"💾 File kết quả: {output_file}")
        
    except Exception as e:
        print(f"❌ Lỗi: {str(e)}")
        return False
    
    return True

def main():
    """Hàm chính"""
    input_file = "kienthucchungfinal.docx"
    output_file = "kienthucchungfinal_cleaned_v2.docx"
    
    # Kiểm tra file đầu vào
    if not os.path.exists(input_file):
        print(f"❌ Không tìm thấy file: {input_file}")
        return
    
    print("🚀 BẮT ĐẦU XÓA TIẾNG VIỆT KHỎI FILE WORD")
    print("=" * 50)
    
    # Xử lý file
    success = remove_vietnamese_from_docx(input_file, output_file)
    
    if success:
        print(f"\n✅ Thành công! File kết quả: {output_file}")
    else:
        print(f"\n❌ Thất bại!")

if __name__ == "__main__":
    main()
