import openai
import json
import re
from docx import Document
from docx.shared import RGBColor
import time
import os
from typing import List, Dict, Optional, Tuple
import hashlib

# Regex patterns for Japanese text detection
JAPANESE_PATTERN = re.compile(r'[\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff]+')

class JapaneseToVietnameseTranslatorChatGPT:
    def __init__(self, api_key: str, model: str = "gpt-3.5-turbo"):
        """Initialize translator with ChatGPT API key"""
        if not api_key or api_key == "your_openai_api_key_here":
            raise ValueError("Please provide a valid OpenAI API key")
        
        self.client = openai.OpenAI(api_key=api_key)
        self.model = model  # gpt-3.5-turbo hoặc gpt-4
        self.translation_cache = {}
        self.request_count = 0
        self.total_tokens_used = 0
        self.estimated_cost = 0.0
        
        # Pricing (USD per 1M tokens) - cập nhật tháng 7/2025
        self.pricing = {
            # Legacy models (converted to per 1K for compatibility)
            "gpt-3.5-turbo": {"input": 0.0005, "output": 0.0015},  # $0.50/$1.50 per 1M
            "gpt-4": {"input": 0.03, "output": 0.06},               # $30/$60 per 1M
            "gpt-4-turbo": {"input": 0.01, "output": 0.03},         # $10/$30 per 1M
            
            # New models (per 1K tokens for compatibility)
            "gpt-4o": {"input": 0.005, "output": 0.015},            # $5/$15 per 1M
            "gpt-4o-mini": {"input": 0.00015, "output": 0.0006},    # $0.15/$0.60 per 1M (RECOMMENDED)
            "gpt-4.1": {"input": 0.002, "output": 0.008},           # $2.00/$8.00 per 1M
            "gpt-4.1-mini": {"input": 0.0004, "output": 0.0016},    # $0.40/$1.60 per 1M
            "gpt-4.1-nano": {"input": 0.0001, "output": 0.0004},    # $0.10/$0.40 per 1M
            "o3": {"input": 0.002, "output": 0.008},                # $2.00/$8.00 per 1M
            "o4-mini": {"input": 0.0011, "output": 0.0044}          # $1.10/$4.40 per 1M
        }
        
        # Batch settings tối ưu cho speed + accuracy
        # Mục tiêu: Output < 32K tokens để tránh limit, đủ context để dịch chính xác
        
        if "gpt-4" in self.model and "4o" not in self.model and "4.1" not in self.model:
            # GPT-4 legacy (8K context limit)
            self.max_batch_size = 6       # Nhỏ vì context hạn chế
            self.max_batch_chars = 3000   # 3K chars ≈ 4.5K tokens input + prompt
            self.batch_delay = 1.5        # Delay để tránh rate limit
            self.max_output_tokens = 2000 # Output limit
            
        elif "gpt-4.1" in self.model:
            # GPT-4.1 series (1M context, 32K output limit)
            # Tối ưu: Input đủ lớn để có context, output < 32K
            self.max_batch_size = 12      # 12 đoạn/batch - cân bằng tốt
            self.max_batch_chars = 8000   # 8K chars ≈ 12K input tokens
            self.batch_delay = 0.5        # Nhanh vì có context lớn
            self.max_output_tokens = 16000 # Output limit: 16K tokens (an toàn)
            
        elif "gpt-4o" in self.model:
            # GPT-4o (128K context)
            self.max_batch_size = 10      # 10 đoạn/batch
            self.max_batch_chars = 6000   # 6K chars để tránh context limit
            self.batch_delay = 0.8
            self.max_output_tokens = 8000 # Output limit
            
        else:
            # GPT-3.5-turbo và model khác
            self.max_batch_size = 8       # 8 đoạn/batch
            self.max_batch_chars = 5000   # 5K chars
            self.batch_delay = 1.0
            self.max_output_tokens = 4000 # Output limit
        
        print(f"🤖 Khởi tạo ChatGPT Translator - Model: {self.model}")
        print(f"📦 Batch settings: {self.max_batch_size} paragraphs, {self.max_batch_chars} chars")
        print(f"🎯 Max output tokens: {self.max_output_tokens}")
        if self.model in self.pricing:
            input_price = self.pricing[self.model]["input"]
            output_price = self.pricing[self.model]["output"]
            print(f"💰 Giá: ${input_price:.4f}/${output_price:.4f} per 1K tokens")
        
    def has_japanese(self, text: str) -> bool:
        """Kiểm tra xem text có chứa ký tự tiếng Nhật không"""
        if not text:
            return False
        return bool(JAPANESE_PATTERN.search(text))
    
    def clean_japanese_text(self, text: str) -> str:
        """Làm sạch text tiếng Nhật"""
        # Loại bỏ whitespace thừa
        text = re.sub(r'\s+', ' ', text).strip()
        
        # Loại bỏ các ký tự đặc biệt không cần thiết
        text = re.sub(r'[^\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff\s.,!?()（）「」『』、。！？]', '', text)
        
        return text
    
    def estimate_tokens(self, text: str) -> int:
        """Ước tính số tokens của text (Japanese text thường ~1.5-2 tokens per character)"""
        # Rough estimation: 1 character ≈ 1.5 tokens cho tiếng Nhật
        return int(len(text) * 1.5)
    
    def calculate_cost(self, input_tokens: int, output_tokens: int) -> float:
        """Tính toán chi phí dựa trên số tokens"""
        if self.model not in self.pricing:
            return 0.0
        
        pricing = self.pricing[self.model]
        input_cost = (input_tokens / 1000) * pricing["input"]
        output_cost = (output_tokens / 1000) * pricing["output"]
        return input_cost + output_cost
    
    def create_translation_batch(self, paragraphs: List[Tuple[str, str]]) -> List[List[Tuple[str, str]]]:
        """Tạo các batch tối ưu từ danh sách paragraphs với kiểm soát output tokens"""
        batches = []
        current_batch = []
        current_chars = 0
        
        for para_id, text in paragraphs:
            cleaned_text = self.clean_japanese_text(text)
            if not cleaned_text or cleaned_text in self.translation_cache:
                continue
            
            text_length = len(cleaned_text)
            
            # Ước tính output tokens cho batch hiện tại + text mới
            estimated_batch_input = current_chars + text_length
            estimated_output_tokens = self.estimate_tokens(str(estimated_batch_input)) // 2  # Output thường ngắn hơn input
            
            # Kiểm tra các giới hạn
            size_ok = len(current_batch) < self.max_batch_size
            chars_ok = estimated_batch_input <= self.max_batch_chars
            output_ok = estimated_output_tokens <= (self.max_output_tokens * 0.8)  # Để 20% buffer
            
            if size_ok and chars_ok and output_ok:
                current_batch.append((para_id, cleaned_text))
                current_chars += text_length
            else:
                # Lưu batch hiện tại và tạo batch mới
                if current_batch:
                    batches.append(current_batch)
                current_batch = [(para_id, cleaned_text)]
                current_chars = text_length
        
        # Thêm batch cuối cùng
        if current_batch:
            batches.append(current_batch)
        
        return batches
    
    def translate_batch(self, batch: List[Tuple[str, str]]) -> Dict[str, str]:
        """Dịch một batch các paragraphs cùng lúc"""
        if not batch:
            return {}
        
        # Tạo unique ID cho mỗi đoạn để tracking
        batch_items = []
        for i, (para_id, text) in enumerate(batch):
            unique_id = f"PARA_{i+1:03d}"
            batch_items.append((unique_id, para_id, text))
        
        # Tạo combined prompt
        japanese_texts = []
        for unique_id, para_id, text in batch_items:
            japanese_texts.append(f"[{unique_id}] {text}")
        
        combined_text = "\n\n".join(japanese_texts)
        
        prompt = f"""Bạn là một chuyên gia dịch thuật tiếng Nhật sang tiếng Việt. Hãy dịch CHÍNH XÁC các đoạn văn sau đây.

QUAN TRỌNG:
- Mỗi đoạn có ID riêng [PARA_XXX]
- Phải trả về ĐÚNG định dạng: [PARA_XXX] Bản dịch tiếng Việt
- Dịch tự nhiên, không máy móc
- Giữ nguyên ngữ cảnh và ý nghĩa
- Dùng từ chuyên ngành cơ khí, kỹ thuật nếu có
- KHÔNG thêm giải thích hay chú thích gì khác

Văn bản tiếng Nhật cần dịch:

{combined_text}

Bản dịch tiếng Việt (giữ đúng format [PARA_XXX]):"""

        try:
            # Ước tính tokens trước khi gọi API
            estimated_input_tokens = self.estimate_tokens(prompt)
            estimated_output_tokens = self.estimate_tokens(combined_text) // 2  # Output thường ngắn hơn
            estimated_cost = self.calculate_cost(estimated_input_tokens, estimated_output_tokens)
            
            print(f"📤 Batch {len(batch)} items - Input: ~{estimated_input_tokens} tokens")
            print(f"   Est Output: ~{estimated_output_tokens} tokens (limit: {self.max_output_tokens})")
            print(f"   Est Cost: ~${estimated_cost:.4f}")
            
            # Kiểm tra nếu estimated output gần limit
            if estimated_output_tokens > self.max_output_tokens * 0.9:
                print(f"   ⚠️ Output gần limit, có thể cần chia nhỏ batch")
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system", 
                        "content": "Bạn là chuyên gia dịch thuật tiếng Nhật sang tiếng Việt chuyên nghiệp, dịch chính xác và tự nhiên."
                    },
                    {
                        "role": "user", 
                        "content": prompt
                    }
                ],
                temperature=0.3,
                max_tokens=self.max_output_tokens,  # Sử dụng dynamic max_tokens
            )
            
            # Cập nhật thống kê
            self.request_count += 1
            if hasattr(response, 'usage'):
                input_tokens = response.usage.prompt_tokens
                output_tokens = response.usage.completion_tokens
                self.total_tokens_used += input_tokens + output_tokens
                actual_cost = self.calculate_cost(input_tokens, output_tokens)
                self.estimated_cost += actual_cost
                
                print(f"✅ Response: {input_tokens} + {output_tokens} tokens = ${actual_cost:.4f}")
            
            # Parse kết quả
            translations = {}
            if response.choices and response.choices[0].message.content:
                content = response.choices[0].message.content.strip()
                
                # Tách theo lines và tìm pattern [PARA_XXX]
                lines = content.split('\n')
                for line in lines:
                    line = line.strip()
                    if line.startswith('[PARA_') and ']' in line:
                        try:
                            end_bracket = line.index(']')
                            unique_id = line[1:end_bracket]
                            translation = line[end_bracket + 1:].strip()
                            
                            # Tìm text gốc tương ứng với unique_id
                            for uid, para_id, original_text in batch_items:
                                if uid == unique_id:
                                    # Lưu vào cache và kết quả
                                    self.translation_cache[original_text] = translation
                                    translations[original_text] = translation
                                    break
                        except (ValueError, IndexError) as e:
                            print(f"⚠️ Parse error for line: {line[:50]}... - {e}")
                            continue
            
            print(f"✅ Batch completed: {len(translations)}/{len(batch)} translations extracted")
            return translations
            
        except Exception as e:
            error_msg = str(e)
            print(f"❌ Batch translation error: {e}")
            
            # Nếu lỗi context length, thử chia nhỏ batch
            if "context_length_exceeded" in error_msg or "maximum context length" in error_msg:
                print(f"🔄 Context limit exceeded, splitting batch into smaller chunks...")
                
                # Chia batch thành 2 phần nhỏ hơn
                half_size = len(batch) // 2
                if half_size > 0:
                    batch1 = batch[:half_size]
                    batch2 = batch[half_size:]
                    
                    print(f"   📦 Splitting into {len(batch1)} + {len(batch2)} items")
                    
                    # Dịch từng phần
                    result1 = self.translate_batch(batch1) if batch1 else {}
                    time.sleep(self.batch_delay)  # Delay giữa các sub-batch
                    result2 = self.translate_batch(batch2) if batch2 else {}
                    
                    # Kết hợp kết quả
                    result1.update(result2)
                    return result1
                else:
                    # Nếu batch chỉ có 1 item, dùng single translation
                    return self._fallback_single_translations(batch)
            
            # Nếu lỗi rate limit, retry với delay lớn hơn
            elif "429" in error_msg or "Too Many Requests" in error_msg:
                print(f"⏳ Rate limit hit, waiting 10 seconds before retry...")
                time.sleep(10)
                return self.translate_batch(batch)  # Retry once
            
            # Các lỗi khác, fallback sang single translation
            else:
                return self._fallback_single_translations(batch)
    
    def _fallback_single_translations(self, batch: List[Tuple[str, str]]) -> Dict[str, str]:
        """Fallback: dịch từng đoạn riêng lẻ khi batch fails"""
        print(f"🔄 Fallback to single translations for {len(batch)} items")
        
        results = {}
        for para_id, text in batch:
            try:
                translation = self.translate_text_single(text)
                if translation and not translation.startswith("[Lỗi dịch"):
                    results[text] = translation
                time.sleep(0.5)  # Delay nhỏ giữa các single requests
            except Exception as e:
                print(f"❌ Single translation failed for {para_id}: {e}")
                continue
        
        return results
    
    def translate_text_single(self, japanese_text: str) -> Optional[str]:
        """Dịch một đoạn text đơn lẻ (fallback method)"""
        if not japanese_text or not self.has_japanese(japanese_text):
            return None
        
        cleaned_text = self.clean_japanese_text(japanese_text)
        if not cleaned_text:
            return None
        
        # Kiểm tra cache
        if cleaned_text in self.translation_cache:
            print(f"✓ Cache hit: {cleaned_text[:50]}...")
            return self.translation_cache[cleaned_text]
        
        try:
            prompt = f"""Dịch đoạn văn tiếng Nhật sau sang tiếng Việt một cách tự nhiên và chính xác:

Yêu cầu:
- Dịch chính xác nghĩa và ngữ cảnh
- Sử dụng từ ngữ tiếng Việt tự nhiên, không máy móc  
- Giữ nguyên cấu trúc câu hợp lý
- Không thêm giải thích hay chú thích
- Chỉ trả về bản dịch tiếng Việt
- Dịch các từ chuyên ngành cơ khí, kỹ thuật nếu có
- Không dịch các ký tự không phải tiếng Nhật

Văn bản tiếng Nhật:
{cleaned_text}

Bản dịch tiếng Việt:"""

            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": "Bạn là chuyên gia dịch thuật tiếng Nhật sang tiếng Việt."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=1000,
            )
            
            if response.choices and response.choices[0].message.content:
                translation = response.choices[0].message.content.strip()
                
                # Cập nhật thống kê
                self.request_count += 1
                if hasattr(response, 'usage'):
                    tokens = response.usage.prompt_tokens + response.usage.completion_tokens
                    self.total_tokens_used += tokens
                    cost = self.calculate_cost(response.usage.prompt_tokens, response.usage.completion_tokens)
                    self.estimated_cost += cost
                
                # Lưu cache
                self.translation_cache[cleaned_text] = translation
                
                print(f"✓ Single translated: {cleaned_text[:50]}... → {translation[:50]}...")
                return translation
            
        except Exception as e:
            print(f"✗ Single translation error: {e}")
            return f"[Lỗi dịch: {str(e)}]"
        
        return None
    
    def add_translation_to_paragraph(self, paragraph, vietnamese_text: str):
        """Thêm bản dịch tiếng Việt vào paragraph"""
        # Thêm xuống dòng
        paragraph.add_run().add_break()
        
        # Thêm text tiếng Việt với format đặc biệt
        vn_run = paragraph.add_run(f"🇻🇳 {vietnamese_text}")
        
        # Định dạng cho text tiếng Việt
        vn_run.font.italic = True
        vn_run.font.color.rgb = RGBColor(0, 100, 0)  # Màu xanh lá đậm
        vn_run.font.size = vn_run.font.size  # Giữ nguyên size
    
    def collect_japanese_paragraphs(self, doc: Document) -> List[Tuple[str, str, object]]:
        """Thu thập tất cả paragraphs tiếng Nhật từ document"""
        japanese_paragraphs = []
        
        # Thu thập từ paragraphs chính
        for i, paragraph in enumerate(doc.paragraphs):
            if self.has_japanese(paragraph.text.strip()):
                para_id = f"main_{i}"
                japanese_paragraphs.append((para_id, paragraph.text.strip(), paragraph))
        
        # Thu thập từ tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        if self.has_japanese(paragraph.text.strip()):
                            para_id = f"table_{table_idx}_{row_idx}_{cell_idx}_{para_idx}"
                            japanese_paragraphs.append((para_id, paragraph.text.strip(), paragraph))
        
        return japanese_paragraphs
    
    def process_word_document(self, input_path: str, output_path: str):
        """Xử lý file Word với batch translation tối ưu"""
        print(f"🔄 Đang xử lý tài liệu: {input_path}")
        start_time = time.time()
        
        try:
            doc = Document(input_path)
            
            # Thu thập tất cả paragraphs tiếng Nhật
            print("📋 Đang thu thập paragraphs tiếng Nhật...")
            japanese_paragraphs = self.collect_japanese_paragraphs(doc)
            
            total_japanese = len(japanese_paragraphs)
            print(f"📝 Tìm thấy {total_japanese} paragraphs tiếng Nhật")
            
            if total_japanese == 0:
                print("ℹ️ Không có text tiếng Nhật nào để dịch")
                return
            
            # Tạo batches tối ưu
            print("🔧 Đang tạo batches tối ưu...")
            paragraph_data = [(para_id, text) for para_id, text, _ in japanese_paragraphs]
            batches = self.create_translation_batch(paragraph_data)
            
            print(f"📦 Tạo được {len(batches)} batches")
            
            # Ước tính chi phí trước khi bắt đầu
            total_chars = sum(len(text) for _, text, _ in japanese_paragraphs)
            estimated_tokens = self.estimate_tokens(str(total_chars))
            estimated_total_cost = self.calculate_cost(estimated_tokens, estimated_tokens // 2)
            print(f"💰 Ước tính chi phí: ~${estimated_total_cost:.4f}")
            
            # Xử lý từng batch
            all_translations = {}
            for batch_idx, batch in enumerate(batches):
                print(f"\n🔄 Processing batch {batch_idx + 1}/{len(batches)} ({len(batch)} items)")
                
                # Dịch batch
                batch_translations = self.translate_batch(batch)
                all_translations.update(batch_translations)
                
                # Delay giữa các batches để tránh rate limit
                if batch_idx < len(batches) - 1:
                    print(f"⏳ Waiting {self.batch_delay}s before next batch...")
                    time.sleep(self.batch_delay)
            
            # Áp dụng translations vào document
            print(f"\n📝 Đang áp dụng {len(all_translations)} bản dịch vào document...")
            translated_count = 0
            
            for para_id, original_text, paragraph_obj in japanese_paragraphs:
                cleaned_text = self.clean_japanese_text(original_text)
                if cleaned_text in all_translations:
                    vietnamese_text = all_translations[cleaned_text]
                    self.add_translation_to_paragraph(paragraph_obj, vietnamese_text)
                    translated_count += 1
                elif cleaned_text in self.translation_cache:
                    vietnamese_text = self.translation_cache[cleaned_text]
                    self.add_translation_to_paragraph(paragraph_obj, vietnamese_text)
                    translated_count += 1
            
            # Lưu document
            print("💾 Đang lưu tài liệu...")
            save_start = time.time()
            doc.save(output_path)
            save_time = time.time() - save_start
            
            # Lưu cache
            cache_file = output_path.replace('.docx', '_translation_cache_chatgpt.json')
            self.save_cache(cache_file)
            
            total_time = time.time() - start_time
            
            # In báo cáo kết quả
            self.print_results(output_path, {
                'total_paragraphs': len(doc.paragraphs),
                'japanese_paragraphs': total_japanese,
                'translated_paragraphs': translated_count,
                'batches_processed': len(batches),
                'cache_size': len(self.translation_cache)
            }, total_time, save_time, cache_file)
            
        except Exception as e:
            print(f"❌ Lỗi xử lý tài liệu: {e}")
            import traceback
            traceback.print_exc()
    
    def save_cache(self, cache_file: str):
        """Lưu cache dịch thuật"""
        try:
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.translation_cache, f, ensure_ascii=False, indent=2)
            print(f"💾 Đã lưu {len(self.translation_cache)} bản dịch vào cache")
        except Exception as e:
            print(f"❌ Lỗi lưu cache: {e}")
    
    def load_cache(self, cache_file: str):
        """Tải cache dịch thuật có sẵn"""
        try:
            if os.path.exists(cache_file):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    self.translation_cache = json.load(f)
                print(f"📥 Đã tải {len(self.translation_cache)} bản dịch từ cache")
            else:
                print("📝 Không tìm thấy cache, bắt đầu từ đầu")
        except Exception as e:
            print(f"❌ Lỗi tải cache: {e}")
    
    def print_results(self, output_path: str, stats: Dict, total_time: float, save_time: float, cache_file: str):
        """In báo cáo kết quả với thông tin chi phí"""
        print("\n" + "="*60)
        print("🎉 KẾT QUẢ DỊCH THUẬT - CHATGPT VERSION")
        print("="*60)
        print(f"📁 File đầu ra: {output_path}")
        print(f"🤖 Model: {self.model}")
        print(f"📊 Tổng paragraphs: {stats['total_paragraphs']}")
        print(f"🇯🇵 Paragraphs tiếng Nhật: {stats['japanese_paragraphs']}")
        print(f"🇻🇳 Paragraphs đã dịch: {stats['translated_paragraphs']}")
        print(f"📦 Batches đã xử lý: {stats['batches_processed']}")
        print(f"💾 Bản dịch trong cache: {stats['cache_size']}")
        print(f"🔄 API calls: {self.request_count}")
        print(f"🎯 Tokens sử dụng: {self.total_tokens_used:,}")
        print(f"💰 Chi phí ước tính: ${self.estimated_cost:.4f}")
        print(f"⏱️  Thời gian xử lý: {total_time:.2f} giây")
        print(f"💾 Thời gian lưu: {save_time:.2f} giây")
        print(f"📂 Cache file: {cache_file}")
        
        if stats['japanese_paragraphs'] > 0:
            success_rate = (stats['translated_paragraphs'] / stats['japanese_paragraphs']) * 100
            print(f"✅ Tỷ lệ thành công: {success_rate:.1f}%")
            
            # Tính hiệu quả batch
            if stats['batches_processed'] > 0:
                avg_per_batch = stats['translated_paragraphs'] / stats['batches_processed']
                print(f"📊 Trung bình: {avg_per_batch:.1f} paragraphs/batch")
        
        print("="*60)

def main():
    """Hàm chính"""
    # Import config từ file riêng
    try:
        from config_chatgpt import get_config, print_config
        CONFIG = get_config()
        if CONFIG is None:
            return
        print_config()
    except ImportError:
        print("❌ Không tìm thấy file config_chatgpt.py!")
        print("   Vui lòng tạo file config_chatgpt.py với OpenAI API key")
        return
    
    print("🌟 JAPANESE TO VIETNAMESE TRANSLATOR - CHATGPT VERSION")
    print("="*60)
    print(f"📥 Input: {CONFIG['input_file']}")
    print(f"📤 Output: {CONFIG['output_file']}")
    print(f"🤖 Model: {CONFIG.get('model', 'gpt-3.5-turbo')}")
    print()
    
    # Kiểm tra file input
    if not os.path.exists(CONFIG['input_file']):
        print(f"❌ Không tìm thấy file input: {CONFIG['input_file']}")
        return
    
    try:
        # Khởi tạo translator
        print("🔧 Đang khởi tạo ChatGPT translator...")
        translator = JapaneseToVietnameseTranslatorChatGPT(
            api_key=CONFIG['api_key'],
            model=CONFIG.get('model', 'gpt-3.5-turbo')
        )
        
        # Tải cache có sẵn
        translator.load_cache(CONFIG['cache_file'])
        
        # Xử lý tài liệu
        translator.process_word_document(CONFIG['input_file'], CONFIG['output_file'])
        
        print("\n🎊 DỊCH THUẬT HOÀN THÀNH!")
        print(f"💰 Tổng chi phí: ${translator.estimated_cost:.4f}")
        
    except Exception as e:
        print(f"❌ Lỗi: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
