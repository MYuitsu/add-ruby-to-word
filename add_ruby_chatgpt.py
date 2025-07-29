import os
import time
import json
import hashlib
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import openai

# === CONFIG ===
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or "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"
CACHE_FILE = "chatgpt_ruby_cache.json"
MODEL = "gpt-4o"  # Hoặc gpt-3.5-turbo nếu bạn muốn tiết kiệm

# === CACHE ===
def load_cache():
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def save_cache(cache):
    with open(CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False, indent=2)

def get_cache_key(text):
    return hashlib.sha256(text.encode("utf-8")).hexdigest()

# === OPENAI ===
def get_ruby_from_chatgpt(text, cache):
    key = get_cache_key(text)
    if key in cache:
        return cache[key]
    prompt = f"""
Bạn là một công cụ phân tích tiếng Nhật. Hãy trả về một object JSON mapping các cụm từ có chứa ký tự Hán tự (Kanji, tức là ký tự trong dải Unicode U+4E00-U+9FFF) trong đoạn sau sang furigana (hiragana), ví dụ:
{{"品質管理": "ひんしつかんり", "顧客": "こきゃく"}}
Chỉ trả về object JSON, không giải thích, không thêm ký tự thừa, không thêm ruby cho từ chỉ có Hiragana/Katakana, không thay thế vào đoạn gốc.

Đoạn cần phân tích:
{text}
"""
    try:
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=2048
        )
        result = response.choices[0].message.content.strip()
        cache[key] = result
        save_cache(cache)
        return result
    except Exception as e:
        print(f"OpenAI API error: {e}")
        return text  # fallback: return original

# === DOCX ===
def replace_paragraph_with_ruby(paragraph, ruby_text):
    # Xóa toàn bộ run cũ
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)

    # Tăng line spacing cho paragraph có tiếng Nhật
    if any(ord(c) >= 0x3040 and ord(c) <= 0x9fff for c in paragraph.text):
        try:
            pPr = paragraph._element.get_or_add_pPr()
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:line'), '380')  # 21pt line spacing (giống add_ruby.py)
            spacing.set(qn('w:lineRule'), 'exact')
            spacing.set(qn('w:before'), '60')   # 3pt space before
            spacing.set(qn('w:after'), '20')    # 1pt space after
        except:
            pass

    # Parse text kiểu Kanji(ふりがな) và chuyển thành XML ruby đúng vị trí
    import re
    def create_ruby_element(kanji, furigana):
        ruby = OxmlElement('w:ruby')
        ruby_pr = OxmlElement('w:rubyPr')
        ruby_align = OxmlElement('w:rubyAlign')
        ruby_align.set(qn('w:val'), 'center')
        ruby_pr.append(ruby_align)
        hgt = OxmlElement('w:hgt')
        hgt.set(qn('w:val'), '32')
        ruby_pr.append(hgt)
        ruby_spacing = OxmlElement('w:lid')
        ruby_spacing.set(qn('w:val'), 'ja-JP')
        ruby_pr.append(ruby_spacing)
        ruby.append(ruby_pr)
        rt = OxmlElement('w:rt')
        rt_r = OxmlElement('w:r')
        rt_rpr = OxmlElement('w:rPr')
        rt_sz = OxmlElement('w:sz')
        rt_sz.set(qn('w:val'), '12')
        rt_rpr.append(rt_sz)
        rt_szcs = OxmlElement('w:szCs')
        rt_szcs.set(qn('w:val'), '12')
        rt_rpr.append(rt_szcs)
        rt_r.append(rt_rpr)
        rt_t = OxmlElement('w:t')
        rt_t.text = furigana
        rt_r.append(rt_t)
        rt.append(rt_r)
        ruby.append(rt)
        ruby_base = OxmlElement('w:rubyBase')
        base_r = OxmlElement('w:r')
        base_rpr = OxmlElement('w:rPr')
        base_sz = OxmlElement('w:sz')
        base_sz.set(qn('w:val'), '22')
        base_rpr.append(base_sz)
        base_szcs = OxmlElement('w:szCs')
        base_szcs.set(qn('w:val'), '22')
        base_rpr.append(base_szcs)
        base_r.append(base_rpr)
        base_t = OxmlElement('w:t')
        base_t.text = kanji
        base_r.append(base_t)
        ruby_base.append(base_r)
        ruby.append(ruby_base)
        return ruby

    # Duyệt text, tìm các cụm Kanji(ふりがな)
    pattern = re.compile(r'(.*?)([^\(\)\n]+)\(([^\(\)\n]+)\)', re.DOTALL)
    text = ruby_text
    pos = 0
    length = len(text)
    while pos < length:
        m = pattern.match(text, pos)
        if m:
            before = m.group(1)
            kanji = m.group(2)
            furigana = m.group(3)
            if before:
                paragraph.add_run(before)
            ruby_elem = create_ruby_element(kanji, furigana)
            paragraph._element.append(ruby_elem)
            pos = m.end()
        else:
            paragraph.add_run(text[pos:])
            break

# === MAIN ===
def process_word_with_chatgpt(input_path, output_path):
    cache = load_cache()
    doc = Document(input_path)
    total = len(doc.paragraphs)
    # Chỉ xử lý 20 paragraph đầu (bao gồm cả ngoài bảng và trong bảng)
    max_paragraphs = 20
    processed_count = 0

    # Xử lý các paragraph ngoài bảng
    for i, paragraph in enumerate(doc.paragraphs):
        if processed_count >= max_paragraphs:
            break
        text = paragraph.text.strip()
        if text and any(ord(c) >= 0x3040 and ord(c) <= 0x9fff for c in text):
            print(f"[{processed_count+1}/{max_paragraphs}] Đang xử lý: {text[:30]}...")
            ruby_text = get_ruby_from_chatgpt(text, cache)
            replace_paragraph_with_ruby(paragraph, ruby_text)
            time.sleep(0.5)
            processed_count += 1

    # Nếu chưa đủ 20, tiếp tục xử lý paragraph trong bảng
    if processed_count < max_paragraphs:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if processed_count >= max_paragraphs:
                            break
                        text = paragraph.text.strip()
                        if text and any(ord(c) >= 0x3040 and ord(c) <= 0x9fff for c in text):
                            print(f"[{processed_count+1}/{max_paragraphs}] Đang xử lý (bảng): {text[:30]}...")
                            ruby_text = get_ruby_from_chatgpt(text, cache)
                            replace_paragraph_with_ruby(paragraph, ruby_text)
                            time.sleep(0.5)
                            processed_count += 1
                    if processed_count >= max_paragraphs:
                        break
                if processed_count >= max_paragraphs:
                    break
            if processed_count >= max_paragraphs:
                break

    doc.save(output_path)
    print(f"Đã lưu file: {output_path}")

if __name__ == "__main__":
    input_file = "KHÔNG FURIGANA - KIẾN THỨC CHUNG.docx"
    output_file = "KHÔNG FURIGANA - KIẾN THỨC CHUNG ruby_chatgpt.docx"
    process_word_with_chatgpt(input_file, output_file)
