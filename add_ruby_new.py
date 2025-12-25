import json
import re
import time
import gc
from collections import defaultdict

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================
# Regex compile (tăng tốc)
# =========================
KANJI_PATTERN = re.compile(r'[\u4e00-\u9fff]')                 # dùng cho "có kanji trong chuỗi"
KANJI_CHAR_PATTERN = re.compile(r'[\u4e00-\u9fff々]')          # dùng cho "một ký tự kanji"
JAPANESE_PATTERN = re.compile(r'[\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff]')

# Set để lưu các kanji không tìm thấy
missing_kanji = set()


def has_kanji(text: str) -> bool:
    """Kiểm tra xem text có chứa ký tự Kanji không"""
    return bool(KANJI_PATTERN.search(text or ""))


def is_kanji_char(ch: str) -> bool:
    """Kiểm tra 1 ký tự có phải kanji không"""
    return bool(KANJI_CHAR_PATTERN.fullmatch(ch or ""))


def has_japanese(text: str) -> bool:
    """Kiểm tra xem text có chứa ký tự tiếng Nhật không"""
    return bool(JAPANESE_PATTERN.search(text or ""))


# =========================================
# Copy formatting (giữ nguyên định dạng run)
# =========================================
def copy_run_rpr(src_run, dst_run):
    """Copy toàn bộ run properties (rPr) từ src_run sang dst_run"""
    try:
        # Xóa rPr cũ nếu có
        if dst_run._element.rPr is not None:
            dst_run._element.remove(dst_run._element.rPr)

        rPr = OxmlElement('w:rPr')
        if src_run is not None and src_run._element.rPr is not None:
            for child in src_run._element.rPr:
                rPr.append(child.__copy__())

        dst_run._element.insert(0, rPr)
    except Exception:
        pass


def copy_run_formatting(source_run, target_rpr):
    """Copy formatting từ source run sang target rPr element, loại bỏ highlighting"""
    if source_run and source_run._element.rPr is not None:
        for child in source_run._element.rPr:
            # Copy tất cả format trừ kích thước font và highlighting
            if child.tag not in [qn('w:sz'), qn('w:szCs'), qn('w:highlight'), qn('w:shd')]:
                try:
                    target_rpr.append(child.__copy__())
                except Exception:
                    pass


# =========================================
# Clean word (radical -> kanji, remove noise)
# =========================================
def clean_kanji_word(word: str) -> str:
    """Làm sạch từ Kanji: thay radical, loại bỏ ký tự không phải tiếng Nhật (giữ dấu câu JP)"""
    radical_to_kanji = {
        "⼀": "一", "⼁": "｜", "⼂": "丶", "⼃": "丿", "⼄": "乙", "⼅": "亅", "⼆": "二", "⼇": "亠", "⼈": "人", "⼉": "儿",
        "⼊": "入", "⼋": "八", "⼌": "冂", "⼍": "冖", "⼎": "冫", "⼏": "几", "⼐": "凵", "⼑": "刀", "⼒": "力", "⼓": "勹",
        "⼔": "匕", "⼕": "匚", "⼖": "匸", "⼗": "十", "⼘": "卜", "⼙": "卩", "⼚": "厂", "⼛": "厶", "⼜": "又", "⼝": "口",
        "⼞": "囗", "⼟": "土", "⼠": "士", "⼡": "夂", "⼢": "夊", "⼣": "夕", "⼤": "大", "⼥": "女", "⼦": "子", "⼧": "宀",
        "⼨": "寸", "⼩": "小", "⼪": "尢", "⼫": "尸", "⼬": "屮", "⼭": "山", "⼮": "巛", "⼯": "工", "⼰": "己", "⼱": "巾",
        "⼲": "干", "⼳": "幺", "⼴": "广", "⼵": "廴", "⼶": "廾", "⼷": "弋", "⼸": "弓", "⼹": "彐", "⼺": "彡", "⼻": "彳",
        "⼼": "心", "⼽": "戈", "⼾": "戸", "⼿": "手", "⽀": "支", "⽁": "攴", "⽂": "文", "⽃": "斗", "⽄": "斤", "⽅": "方",
        "⽆": "无", "⽇": "日", "⽈": "曰", "⽉": "月", "⽊": "木", "⽋": "欠", "⽌": "止", "⽍": "歹", "⽎": "殳", "⽏": "毋",
        "⽐": "比", "⽑": "毛", "⽒": "氏", "⽓": "气", "⽔": "水", "⽕": "火", "⽖": "爪", "⽗": "父", "⽘": "爻", "⽙": "爿",
        "⽚": "片", "⽛": "牙", "⽜": "牛", "⽝": "犬", "⽞": "玄", "⽟": "玉", "⽠": "瓜", "⽡": "瓦", "⽢": "甘", "⽣": "生",
        "⽤": "用", "⽥": "田", "⽦": "疋", "⽧": "疒", "⽨": "癶", "⽩": "白", "⽪": "皮", "⽫": "皿", "⽬": "目", "⽭": "矛",
        "⽮": "矢", "⽯": "石", "⽰": "示", "⽱": "禸", "⽲": "禾", "⽳": "穴", "⽴": "立", "⽵": "竹", "⽶": "米", "⽷": "糸",
        "⽸": "缶", "⽹": "网", "⽺": "羊", "⽻": "羽", "⽼": "老", "⽽": "而", "⽾": "耒", "⽿": "耳", "⾀": "聿", "⾁": "肉",
        "⾂": "臣", "⾃": "自", "⾄": "至", "⾅": "臼", "⾆": "舌", "⾇": "舛", "⾈": "舟", "⾉": "艮", "⾊": "色", "⾋": "艸",
        "⾌": "虍", "⾍": "虫", "⾎": "血", "⾏": "行", "⾐": "衣", "⾑": "襾", "⾒": "見", "⾓": "角", "⾔": "言", "⾕": "谷",
        "⾖": "豆", "⾗": "豕", "⾘": "豸", "⾙": "貝", "⾚": "赤", "⾛": "走", "⾜": "足", "⾝": "身", "⾞": "車", "⾟": "辛",
        "⾠": "辰", "⾡": "辵", "⾢": "邑", "⾣": "酉", "⾤": "釆", "⾥": "里", "⾦": "金", "⾧": "長", "⾨": "門", "⾩": "阜",
        "⾪": "隶", "⾫": "隹", "⾬": "雨", "⾭": "青", "⾮": "非", "⾯": "面", "⾰": "革", "⾱": "韋", "⾲": "韭", "⾳": "音",
        "⾴": "頁", "⾵": "風", "⾶": "飛", "⾷": "食", "⾸": "首", "⾹": "香", "⾺": "馬", "⾻": "骨", "⾼": "高", "⾽": "髟",
        "⾾": "鬥", "⾿": "鬯", "⿀": "鬲", "⿁": "鬼", "⿂": "魚", "⿃": "鳥", "⿄": "鹵", "⿅": "鹿", "⿆": "麦", "⿇": "麻",
        "⿈": "黄", "⿉": "黍", "⿊": "黒", "⿋": "黹", "⿌": "黽", "⿍": "鼎", "⿎": "鼓", "⿏": "鼠", "⿐": "鼻", "⿑": "齊",
        "⿒": "歯", "⿓": "竜", "⿔": "亀", "⿕": "龠"
    }

    word = ''.join([radical_to_kanji.get(ch, ch) for ch in (word or "")])

    cleaned = re.sub(
        r'[^\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff0-9０-９、。！？「」『』（）［］【】〈〉《》〔〕…‥ー～—・：；]',
        '',
        word
    )
    return cleaned


# =========================================
# Dictionary schema mới/cũ -> normalize
# =========================================
def _build_map_from_segments(surface: str, segments):
    """
    Nếu dict có 'segments': [{"s":[s0,s1],"rt":"..."}], build map đơn giản.
    (Khuyến nghị: nếu muốn đúng tuyệt đối per-kanji thì nên dùng map chuẩn thay vì segments.)
    """
    mp = []
    if not isinstance(segments, list):
        return None
    for seg in segments:
        try:
            s0, s1 = seg["s"]
            rt = seg["rt"]
            for i in range(s0, s1):
                if 0 <= i < len(surface) and is_kanji_char(surface[i]):
                    mp.append({"i": i, "ch": surface[i], "rt": rt})
        except Exception:
            continue
    return mp or None


def _normalize_dict_value(key: str, val):
    """
    Trả về entry chuẩn: {"rt": "...", "map": [...] or None}
    - hỗ trợ schema cũ: "key": "hiragana"
    - hỗ trợ schema mới: "key": {"rt": "...", "map":[...]} (và optionally segments)
    """
    if isinstance(val, str):
        rt = val.strip()
        if rt:
            return {"rt": rt, "map": None}
        return None

    if isinstance(val, dict):
        rt = val.get("rt") or val.get("reading") or val.get("hiragana")
        if isinstance(rt, str):
            rt = rt.strip()
        if not rt:
            return None

        mp = val.get("map")

        # nếu chưa có map nhưng có segments -> build map đơn giản
        if mp is None and "segments" in val:
            mp = _build_map_from_segments(key, val.get("segments"))

        # validate map (i phải là int trong range)
        if isinstance(mp, list):
            ok = True
            for m in mp:
                if not isinstance(m, dict):
                    ok = False
                    break
                i = m.get("i")
                if not isinstance(i, int) or i < 0 or i >= len(key):
                    ok = False
                    break
                if not isinstance(m.get("rt"), str):
                    ok = False
                    break
            if not ok:
                mp = None

        return {"rt": rt, "map": mp}

    return None


def load_dictionary(dictionary_path):
    """Đọc dictionary từ file JSON và tối ưu hóa cho tìm kiếm nhanh (schema mới/cũ)"""
    print("Đang load dictionary...")
    start_time = time.time()

    try:
        with open(dictionary_path, 'r', encoding='utf-8') as f:
            dictionary = json.load(f)

        optimized_dict = {}
        kanji_count = 0

        for key, val in dictionary.items():
            if (has_kanji(key) and
                1 <= len(key) <= 12 and
                not re.search(r'[a-zA-Z0-9]', key)):

                norm = _normalize_dict_value(key, val)
                if norm is None:
                    continue

                optimized_dict[key] = norm
                kanji_count += 1

        sorted_dict = dict(sorted(optimized_dict.items(), key=lambda x: len(x[0]), reverse=True))

        load_time = time.time() - start_time
        print(f"Đã load {kanji_count} từ Kanji trong {load_time:.2f} giây")
        return sorted_dict

    except FileNotFoundError:
        print(f"Không tìm thấy file dictionary: {dictionary_path}")
        return {}
    except json.JSONDecodeError:
        print(f"Lỗi định dạng JSON trong file: {dictionary_path}")
        return {}


# =========================================
# Ruby element builder (Word XML)
# =========================================
def create_ruby_element(kanji, hiragana, source_run=None):
    """Tạo element ruby cho Word document với spacing tối ưu để tránh che khuất"""
    ruby = OxmlElement('w:ruby')

    ruby_pr = OxmlElement('w:rubyPr')
    ruby_align = OxmlElement('w:rubyAlign')
    ruby_align.set(qn('w:val'), 'center')
    ruby_pr.append(ruby_align)

    # khoảng cách ruby/kanji
    hgt = OxmlElement('w:hgt')
    hgt.set(qn('w:val'), '32')
    ruby_pr.append(hgt)

    ruby_spacing = OxmlElement('w:lid')
    ruby_spacing.set(qn('w:val'), 'ja-JP')
    ruby_pr.append(ruby_spacing)

    ruby.append(ruby_pr)

    # Ruby text (hiragana)
    rt = OxmlElement('w:rt')
    rt_r = OxmlElement('w:r')
    rt_rpr = OxmlElement('w:rPr')

    copy_run_formatting(source_run, rt_rpr)

    rt_sz = OxmlElement('w:sz')
    rt_sz.set(qn('w:val'), '12')  # 6pt
    rt_rpr.append(rt_sz)

    rt_szcs = OxmlElement('w:szCs')
    rt_szcs.set(qn('w:val'), '12')
    rt_rpr.append(rt_szcs)

    rt_r.append(rt_rpr)
    rt_t = OxmlElement('w:t')
    if ' ' in hiragana:
        rt_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    rt_t.text = hiragana
    rt_r.append(rt_t)
    rt.append(rt_r)
    ruby.append(rt)

    # Ruby base (kanji)
    ruby_base = OxmlElement('w:rubyBase')
    base_r = OxmlElement('w:r')
    base_rpr = OxmlElement('w:rPr')

    copy_run_formatting(source_run, base_rpr)

    base_sz = OxmlElement('w:sz')
    base_sz.set(qn('w:val'), '22')  # 11pt
    base_rpr.append(base_sz)

    base_szcs = OxmlElement('w:szCs')
    base_szcs.set(qn('w:val'), '22')
    base_rpr.append(base_szcs)

    base_r.append(base_rpr)
    base_t = OxmlElement('w:t')
    base_t.text = kanji
    base_r.append(base_t)
    ruby_base.append(base_r)
    ruby.append(ruby_base)

    return ruby


def append_surface_with_map(paragraph, surface, mp, src_run):
    """
    Apply ruby theo 'map' per-char:
      mp: [{"i":0,"ch":"取","rt":"と"}, ...] với i là index trong surface
    """
    idx_to_rt = {}
    for m in mp or []:
        try:
            i = m["i"]
            rt = m["rt"]
            if isinstance(i, int) and isinstance(rt, str):
                idx_to_rt[i] = rt
        except Exception:
            continue

    for idx, ch in enumerate(surface):
        if idx in idx_to_rt and is_kanji_char(ch):
            paragraph._element.append(create_ruby_element(ch, idx_to_rt[idx], src_run))
        else:
            r = paragraph.add_run(ch)
            copy_run_rpr(src_run, r)


# =========================================
# Match finder: return (start, end, key, rt, map)
# =========================================
def find_kanji_matches_optimized(text, dictionary):
    """Tìm matches; mỗi match trả (start, end, surface_key, rt, map_or_none)"""
    if not text or not has_japanese(text):
        return []

    matches = []
    text_len = len(text)
    covered = [False] * text_len

    def try_add_match(i, length):
        if any(covered[i:i + length]):
            return

        substring = text[i:i + length]
        cleaned_substring = clean_kanji_word(substring)

        if not cleaned_substring or not has_kanji(cleaned_substring):
            return

        if cleaned_substring in dictionary:
            entry = dictionary[cleaned_substring]
            rt = entry.get("rt") if isinstance(entry, dict) else entry
            mp = entry.get("map") if isinstance(entry, dict) else None

            matches.append((i, i + length, cleaned_substring, rt, mp))

            for j in range(i, i + length):
                covered[j] = True

    # Pass 1: từ dài -> ngắn
    for length in range(min(12, text_len), 2, -1):
        for i in range(text_len - length + 1):
            try_add_match(i, length)

    # Pass 2: 2 -> 1
    for length in range(2, 0, -1):
        for i in range(text_len - length + 1):
            try_add_match(i, length)

    # Ghi lại Kanji không tìm thấy
    for i in range(text_len):
        if not covered[i] and has_kanji(text[i]):
            missing_kanji.add(text[i])

    matches.sort(key=lambda x: x[0])
    return matches


# =========================================
# Section break helpers
# =========================================
def has_section_break(paragraph):
    """Kiểm tra xem paragraph có chứa section break không"""
    pPr = paragraph._element.pPr
    if pPr is not None:
        sectPr = pPr.find(qn('w:sectPr'))
        return sectPr is not None
    return False


def remove_section_breaks_and_add_spacing(paragraph):
    """Xóa section break (việc chèn 2 dòng trống sẽ xử lý ở process_word_document)"""
    pPr = paragraph._element.pPr
    if pPr is not None:
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            pPr.remove(sectPr)
            print("Đã xóa section break và thêm 2 dòng trống")
            return True
    return False


# =========================================
# Core: add ruby to paragraph preserving runs
# =========================================
def add_ruby_to_paragraph_preserve_runs(paragraph, dictionary):
    """Thêm ruby vào paragraph, giữ nguyên định dạng/run, hỗ trợ dict schema mới (map)."""

    # Line spacing cho paragraph có tiếng Nhật
    if has_japanese(paragraph.text):
        try:
            pPr = paragraph._element.get_or_add_pPr()
            spacing = pPr.find(qn('w:spacing'))
            if spacing is None:
                spacing = OxmlElement('w:spacing')
                pPr.append(spacing)
            spacing.set(qn('w:line'), '380')  # 21pt line spacing
            spacing.set(qn('w:lineRule'), 'exact')
            spacing.set(qn('w:before'), '60')  # 3pt before
            spacing.set(qn('w:after'), '20')   # 1pt after
        except Exception:
            pass

    section_break_removed = remove_section_breaks_and_add_spacing(paragraph)
    if not paragraph.runs or not has_japanese(paragraph.text):
        return section_break_removed

    # Lưu run gốc
    old_runs = list(paragraph.runs)

    # Xóa run gốc
    for run in old_runs:
        run._element.getparent().remove(run._element)

    for run in old_runs:
        text = run.text or ""

        if not has_japanese(text):
            new_run = paragraph.add_run(text)
            copy_run_rpr(run, new_run)
            continue

        matches = find_kanji_matches_optimized(text, dictionary)
        if not matches:
            new_run = paragraph.add_run(text)
            copy_run_rpr(run, new_run)
            continue

        last_pos = 0
        for start, end, kanji_key, hiragana, mp in matches:
            # text trước match
            if start > last_pos:
                before_text = text[last_pos:start]
                if before_text:
                    new_run = paragraph.add_run(before_text)
                    copy_run_rpr(run, new_run)

            # ✅ Nếu có map -> apply per-character, bỏ heuristics cũ
            if mp:
                append_surface_with_map(paragraph, kanji_key, mp, run)
                last_pos = end
                continue

            # Fallback (schema cũ / không có map): heuristics như bạn đang dùng
            hiragana_actual_len = len((hiragana or "").strip())

            if hiragana_actual_len < len(kanji_key) and hiragana_actual_len > 0:
                kanji_ruby = kanji_key[:hiragana_actual_len]
                ruby_element = create_ruby_element(kanji_ruby, hiragana, run)
                paragraph._element.append(ruby_element)

                if len(kanji_key) > hiragana_actual_len:
                    new_run = paragraph.add_run(kanji_key[hiragana_actual_len:])
                    copy_run_rpr(run, new_run)

                last_pos = end

            elif len(kanji_key) > 1 and hiragana_actual_len == 1:
                # chỉ add ruby cho kanji đầu tiên trong cụm
                for idx, ch in enumerate(kanji_key):
                    if has_kanji(ch):
                        if idx > 0:
                            new_run = paragraph.add_run(kanji_key[:idx])
                            copy_run_rpr(run, new_run)

                        ruby_element = create_ruby_element(ch, hiragana, run)
                        paragraph._element.append(ruby_element)

                        if idx + 1 < len(kanji_key):
                            new_run = paragraph.add_run(kanji_key[idx + 1:])
                            copy_run_rpr(run, new_run)
                        break
                last_pos = end

            else:
                ruby_element = create_ruby_element(kanji_key, hiragana, run)
                paragraph._element.append(ruby_element)
                last_pos = end

        # text còn lại sau cùng
        if last_pos < len(text):
            after_text = text[last_pos:]
            if after_text:
                new_run = paragraph.add_run(after_text)
                copy_run_rpr(run, new_run)

    return True


# =========================================
# Missing report
# =========================================
def save_missing_kanji_report(missing_kanji, output_path):
    """Lưu báo cáo các từ Kanji không tìm thấy"""
    if not missing_kanji:
        print("Tất cả từ Kanji đều có trong dictionary!")
        return

    report_file = output_path.replace('.docx', '_missing_kanji.txt')
    sorted_missing = sorted(missing_kanji, key=lambda x: (len(x), x))

    with open(report_file, 'w', encoding='utf-8') as f:
        f.write("=== BÁO CÁO CÁC TỪ KANJI KHÔNG TÌM THẤY TRONG DICTIONARY ===\n\n")
        f.write(f"Tổng số từ Kanji không tìm thấy: {len(missing_kanji)}\n\n")

        by_length = defaultdict(list)
        for word in sorted_missing:
            by_length[len(word)].append(word)

        for length in sorted(by_length.keys()):
            f.write(f"=== Từ có {length} ký tự ({len(by_length[length])} từ) ===\n")
            for word in by_length[length]:
                f.write(f"{word}\n")
            f.write("\n")

        f.write("=== DANH SÁCH ĐẦY ĐỦ ===\n")
        for word in sorted_missing:
            f.write(f"{word}\n")

    print(f"Đã lưu báo cáo từ Kanji thiếu: {report_file}")
    print(f"Tổng số từ Kanji không tìm thấy: {len(missing_kanji)}")


# =========================================
# Process Word document
# =========================================
def process_word_document(input_path, output_path, dictionary_path):
    """Xử lý file Word"""
    global missing_kanji
    missing_kanji.clear()

    print("=== Xử lý file Word (hỗ trợ dictionary schema mới) ===")
    start_time = time.time()

    dictionary = load_dictionary(dictionary_path)
    if not dictionary:
        print("Dictionary trống hoặc không đọc được.")
        return

    print(f"Đang xử lý file: {input_path}")

    try:
        doc = Document(input_path)

        japanese_paragraphs = 0
        total_paragraphs = 0
        processed_count = 0
        section_breaks_removed = 0

        print("Đang xử lý paragraphs chính...")

        paragraphs_to_add_spacing = []

        for i, paragraph in enumerate(doc.paragraphs):
            total_paragraphs += 1

            if has_section_break(paragraph):
                section_breaks_removed += 1
                paragraphs_to_add_spacing.append(i)

            if paragraph.text.strip() and has_japanese(paragraph.text):
                japanese_paragraphs += 1

                if add_ruby_to_paragraph_preserve_runs(paragraph, dictionary):
                    processed_count += 1

                if japanese_paragraphs % 50 == 0:
                    elapsed = time.time() - start_time
                    print(f"Đã xử lý {japanese_paragraphs} paragraph tiếng Nhật - {elapsed:.1f}s")

                    if japanese_paragraphs % 200 == 0:
                        gc.collect()
            else:
                add_ruby_to_paragraph_preserve_runs(paragraph, dictionary)

        # Thêm 2 dòng trống sau các paragraph có section break
        for i in reversed(paragraphs_to_add_spacing):
            p1 = doc.add_paragraph()
            p2 = doc.add_paragraph()
            doc._element.body.insert(i + 1, p1._element)
            doc._element.body.insert(i + 2, p2._element)

        print("Đang xử lý tables...")

        table_count = 0
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total_paragraphs += 1
                        if paragraph.text.strip() and has_japanese(paragraph.text):
                            japanese_paragraphs += 1
                            if add_ruby_to_paragraph_preserve_runs(paragraph, dictionary):
                                processed_count += 1

            if table_count % 20 == 0:
                elapsed = time.time() - start_time
                print(f"Đã xử lý {table_count} tables - {elapsed:.1f}s")

        print("Đang lưu file...")
        save_start = time.time()
        doc.save(output_path)
        save_time = time.time() - save_start

        save_missing_kanji_report(missing_kanji, output_path)

        total_time = time.time() - start_time

        print("\n=== KẾT QUẢ XỬ LÝ ===")
        print(f"File output: {output_path}")
        print(f"Tổng số paragraph: {total_paragraphs}")
        print(f"Paragraph có tiếng Nhật: {japanese_paragraphs}")
        print(f"Paragraph đã thêm ruby: {processed_count}")
        if japanese_paragraphs > 0:
            print(f"Tỷ lệ xử lý thành công: {processed_count / japanese_paragraphs * 100:.1f}%")
        print(f"Section breaks đã xóa: {section_breaks_removed}")
        print(f"Từ Kanji không tìm thấy: {len(missing_kanji)}")
        print(f"Thời gian xử lý: {total_time:.2f} giây")
        print(f"Thời gian lưu file: {save_time:.2f} giây")
        if japanese_paragraphs > 0:
            print(f"Tốc độ xử lý: {japanese_paragraphs / total_time:.1f} paragraph/giây")

    except Exception as e:
        print(f"Lỗi khi xử lý file Word: {e}")
        import traceback
        traceback.print_exc()


# =========================================
# Optional helper (giữ lại như file cũ)
# =========================================
def get_first_run_from_paragraph(paragraph):
    """Lấy run đầu tiên từ paragraph để copy format"""
    if paragraph.runs:
        return paragraph.runs[0]
    return None


def print_processing_stats(japanese_paragraphs, processed_count, missing_count, total_time):
    """In thống kê quá trình xử lý"""
    print(f"\n=== THỐNG KÊ XỬ LÝ ===")
    print(f"Paragraph có tiếng Nhật: {japanese_paragraphs}")
    print(f"Paragraph đã thêm ruby: {processed_count}")
    if japanese_paragraphs > 0:
        print(f"Tỷ lệ xử lý: {processed_count / japanese_paragraphs * 100:.1f}%")
    print(f"Từ Kanji không tìm thấy: {missing_count}")
    print(f"Thời gian xử lý: {total_time:.2f} giây")
    if japanese_paragraphs > 0:
        print(f"Tốc độ: {japanese_paragraphs / total_time:.1f} paragraph/giây")


def create_japanese_run_with_font_size(paragraph, text, font_size_pt=11):
    """Tạo run với kích thước font cụ thể cho text tiếng Nhật"""
    run = paragraph.add_run(text)

    rPr = run._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)

    for sz in rPr.findall(qn('w:sz')):
        rPr.remove(sz)
    for szcs in rPr.findall(qn('w:szCs')):
        rPr.remove(szcs)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size_pt * 2))
    rPr.append(sz)

    szcs = OxmlElement('w:szCs')
    szcs.set(qn('w:val'), str(font_size_pt * 2))
    rPr.append(szcs)

    return run


# =========================================
# Main
# =========================================
def main():
    input_file = "Operetion 2025.03.24.docx"
    output_file = "asdasdasd_ruby_test.docx"
    dictionary_file = "dict_struct.json"  # giờ có thể là schema mới/cũ đều được

    print("=== Chương trình thêm Ruby cho file Word (schema dictionary mới) ===")
    print(f"Input file: {input_file}")
    print(f"Output file: {output_file}")
    print(f"Dictionary file: {dictionary_file}")
    print()

    process_word_document(input_file, output_file, dictionary_file)
    print("\nHoàn thành!")


if __name__ == "__main__":
    main()
