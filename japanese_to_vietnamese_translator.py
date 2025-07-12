import google.generativeai as genai
import json
import re
from docx import Document
from docx.shared import Inches
import time
import os
from typing import List, Tuple

# Configure Gemini API
GEMINI_API_KEY = "your_gemini_api_key_here"  # Replace with your actual API key
genai.configure(api_key=GEMINI_API_KEY)

# Initialize Gemini model
model = genai.GenerativeModel('gemini-pro')

# Regex patterns for Japanese text detection
JAPANESE_PATTERN = re.compile(r'[\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff]+')

class JapaneseTranslator:
    def __init__(self, api_key: str = None):
        if api_key:
            genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-pro')
        self.translation_cache = {}
    
    def has_japanese(self, text: str) -> bool:
        """Check if text contains Japanese characters"""
        return bool(JAPANESE_PATTERN.search(text))
    
    def extract_japanese_sentences(self, text: str) -> List[str]:
        """Extract Japanese sentences from text"""
        if not self.has_japanese(text):
            return []
        
        # Split by common Japanese sentence endings
        sentences = re.split(r'[。！？\n]', text)
        japanese_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and self.has_japanese(sentence):
                japanese_sentences.append(sentence)
        
        return japanese_sentences
    
    def translate_text(self, japanese_text: str) -> str:
        """Translate Japanese text to Vietnamese using Gemini API"""
        if not japanese_text or not self.has_japanese(japanese_text):
            return ""
        
        # Check cache first
        if japanese_text in self.translation_cache:
            return self.translation_cache[japanese_text]
        
        try:
            prompt = f"""
            Translate the following Japanese text to Vietnamese accurately and naturally.
            Keep the meaning precise and use appropriate Vietnamese expressions.
            Only return the Vietnamese translation, no explanations.
            
            Japanese text: {japanese_text}
            """
            
            response = self.model.generate_content(prompt)
            vietnamese_translation = response.text.strip()
            
            # Cache the translation
            self.translation_cache[japanese_text] = vietnamese_translation
            
            # Add delay to respect API rate limits
            time.sleep(0.1)
            
            return vietnamese_translation
            
        except Exception as e:
            print(f"Error translating '{japanese_text}': {e}")
            return f"[Translation error: {str(e)}]"
    
    def process_paragraph(self, paragraph) -> bool:
        """Process a paragraph and add Vietnamese translation below Japanese text"""
        original_text = paragraph.text.strip()
        
        if not original_text or not self.has_japanese(original_text):
            return False
        
        # Extract Japanese sentences
        japanese_sentences = self.extract_japanese_sentences(original_text)
        
        if not japanese_sentences:
            return False
        
        # Translate each sentence
        translations = []
        for sentence in japanese_sentences:
            translation = self.translate_text(sentence)
            if translation:
                translations.append(translation)
        
        if not translations:
            return False
        
        # Combine translations
        vietnamese_text = " ".join(translations)
        
        # Clear existing runs
        for run in paragraph.runs:
            run._element.getparent().remove(run._element)
        
        # Add original Japanese text
        paragraph.add_run(original_text)
        
        # Add line break
        paragraph.add_run("\n")
        
        # Add Vietnamese translation with different formatting
        vietnamese_run = paragraph.add_run(f"[VN]: {vietnamese_text}")
        vietnamese_run.font.italic = True
        vietnamese_run.font.color.rgb = None  # Use default color
        
        return True
    
    def process_word_document(self, input_path: str, output_path: str):
        """Process Word document and add Vietnamese translations"""
        print(f"Processing document: {input_path}")
        start_time = time.time()
        
        try:
            doc = Document(input_path)
            
            total_paragraphs = 0
            japanese_paragraphs = 0
            translated_paragraphs = 0
            
            print("Processing main paragraphs...")
            
            # Process main document paragraphs
            for i, paragraph in enumerate(doc.paragraphs):
                total_paragraphs += 1
                
                if self.has_japanese(paragraph.text):
                    japanese_paragraphs += 1
                    
                    if self.process_paragraph(paragraph):
                        translated_paragraphs += 1
                    
                    # Show progress every 20 paragraphs
                    if japanese_paragraphs % 20 == 0:
                        elapsed = time.time() - start_time
                        print(f"Processed {japanese_paragraphs} Japanese paragraphs - {elapsed:.1f}s")
            
            print("Processing tables...")
            
            # Process tables
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            total_paragraphs += 1
                            
                            if self.has_japanese(paragraph.text):
                                japanese_paragraphs += 1
                                
                                if self.process_paragraph(paragraph):
                                    translated_paragraphs += 1
                
                # Show progress every 10 tables
                if table_count % 10 == 0:
                    elapsed = time.time() - start_time
                    print(f"Processed {table_count} tables - {elapsed:.1f}s")
            
            # Save the document
            print("Saving document...")
            save_start = time.time()
            doc.save(output_path)
            save_time = time.time() - save_start
            
            # Save translation cache
            cache_file = output_path.replace('.docx', '_translation_cache.json')
            with open(cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.translation_cache, f, ensure_ascii=False, indent=2)
            
            total_time = time.time() - start_time
            
            print("\n=== Translation Results ===")
            print(f"Output file: {output_path}")
            print(f"Total paragraphs: {total_paragraphs}")
            print(f"Japanese paragraphs: {japanese_paragraphs}")
            print(f"Translated paragraphs: {translated_paragraphs}")
            print(f"Translations cached: {len(self.translation_cache)}")
            print(f"Processing time: {total_time:.2f} seconds")
            print(f"Save time: {save_time:.2f} seconds")
            print(f"Translation cache saved: {cache_file}")
            
        except Exception as e:
            print(f"Error processing document: {e}")
            import traceback
            traceback.print_exc()
    
    def load_translation_cache(self, cache_file: str):
        """Load existing translation cache to avoid re-translating"""
        try:
            if os.path.exists(cache_file):
                with open(cache_file, 'r', encoding='utf-8') as f:
                    self.translation_cache = json.load(f)
                print(f"Loaded {len(self.translation_cache)} cached translations")
        except Exception as e:
            print(f"Error loading cache: {e}")

def main():
    # Configuration
    input_file = "Operetion 2025.03.24.docx"
    output_file = "Operetion 2025.03.24_with_vietnamese.docx"
    api_key = "your_gemini_api_key_here"  # Replace with your actual API key
    cache_file = "translation_cache.json"
    
    print("=== Japanese to Vietnamese Translator ===")
    print(f"Input file: {input_file}")
    print(f"Output file: {output_file}")
    print()
    
    # Initialize translator
    translator = JapaneseTranslator(api_key)
    
    # Load existing cache if available
    translator.load_translation_cache(cache_file)
    
    # Process the document
    translator.process_word_document(input_file, output_file)
    
    print("\nTranslation completed!")

if __name__ == "__main__":
    main()