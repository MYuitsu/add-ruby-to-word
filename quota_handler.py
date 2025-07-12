#!/usr/bin/env python3
"""
Script xử lý khi gặp daily quota limit - tạo partial translation và kế hoạch tiếp tục
"""

import json
import os
from docx import Document
import re
from typing import List, Dict

# Regex patterns for Japanese text detection
JAPANESE_PATTERN = re.compile(r'[\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff]+')

def has_japanese(text: str) -> bool:
    """Kiểm tra xem text có chứa ký tự tiếng Nhật không"""
    if not text:
        return False
    return bool(JAPANESE_PATTERN.search(text))

def analyze_remaining_work(input_file: str, cache_file: str):
    """Phân tích công việc còn lại sau khi vượt quota"""
    print("📊 PHÂN TÍCH CÔNG VIỆC CÒN LẠI")
    print("=" * 50)
    
    # Đọc cache hiện tại
    translation_cache = {}
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'r', encoding='utf-8') as f:
                translation_cache = json.load(f)
            print(f"📥 Đã có {len(translation_cache)} bản dịch trong cache")
        except Exception as e:
            print(f"❌ Lỗi đọc cache: {e}")
    else:
        print("📝 Chưa có cache nào")
    
    # Phân tích file Word
    try:
        doc = Document(input_file)
        
        stats = {
            'total_paragraphs': 0,
            'japanese_paragraphs': 0,
            'already_translated': 0,
            'remaining_to_translate': 0,
            'japanese_texts': []
        }
        
        print("🔍 Đang phân tích paragraphs...")
        
        # Phân tích paragraphs chính
        for paragraph in doc.paragraphs:
            stats['total_paragraphs'] += 1
            text = paragraph.text.strip()
            
            if text and has_japanese(text):
                stats['japanese_paragraphs'] += 1
                
                # Làm sạch text như trong translator
                cleaned_text = re.sub(r'\s+', ' ', text).strip()
                cleaned_text = re.sub(r'[^\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff\s.,!?()（）「」『』、。！？]', '', cleaned_text)
                
                if cleaned_text in translation_cache:
                    stats['already_translated'] += 1
                else:
                    stats['remaining_to_translate'] += 1
                    stats['japanese_texts'].append(cleaned_text)
        
        # Phân tích tables
        print("🔍 Đang phân tích tables...")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        stats['total_paragraphs'] += 1
                        text = paragraph.text.strip()
                        
                        if text and has_japanese(text):
                            stats['japanese_paragraphs'] += 1
                            
                            cleaned_text = re.sub(r'\s+', ' ', text).strip()
                            cleaned_text = re.sub(r'[^\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff\s.,!?()（）「」『』、。！？]', '', cleaned_text)
                            
                            if cleaned_text in translation_cache:
                                stats['already_translated'] += 1
                            else:
                                stats['remaining_to_translate'] += 1
                                if cleaned_text not in stats['japanese_texts']:
                                    stats['japanese_texts'].append(cleaned_text)
        
        # In báo cáo
        print("\n📈 KẾT QUẢ PHÂN TÍCH:")
        print("-" * 40)
        print(f"📄 Tổng paragraphs: {stats['total_paragraphs']}")
        print(f"🇯🇵 Paragraphs tiếng Nhật: {stats['japanese_paragraphs']}")
        print(f"✅ Đã dịch: {stats['already_translated']}")
        print(f"⏳ Còn lại: {stats['remaining_to_translate']}")
        
        if stats['japanese_paragraphs'] > 0:
            progress = (stats['already_translated'] / stats['japanese_paragraphs']) * 100
            print(f"📊 Tiến độ: {progress:.1f}%")
        
        # Tính thời gian cần thiết
        days_needed = (stats['remaining_to_translate'] / 45) + 1  # 45 requests/day để an toàn
        print(f"⏰ Số ngày cần để hoàn thành: {days_needed:.1f}")
        
        # Lưu danh sách văn bản chưa dịch
        remaining_file = input_file.replace('.docx', '_remaining_texts.json')
        with open(remaining_file, 'w', encoding='utf-8') as f:
            json.dump({
                'stats': stats,
                'remaining_texts': stats['japanese_texts'][:stats['remaining_to_translate']]  # Loại bỏ duplicates
            }, f, ensure_ascii=False, indent=2)
        
        print(f"💾 Đã lưu danh sách văn bản chưa dịch: {remaining_file}")
        
        return stats
        
    except Exception as e:
        print(f"❌ Lỗi phân tích file: {e}")
        return None

def create_daily_plan(stats: Dict, max_daily_requests: int = 45):
    """Tạo kế hoạch dịch thuật hàng ngày"""
    if not stats or stats['remaining_to_translate'] <= 0:
        print("🎉 Không có gì để dịch thêm!")
        return
    
    print(f"\n📅 KẾ HOẠCH DỊCH THUẬT HÀNG NGÀY")
    print("=" * 50)
    
    remaining = stats['remaining_to_translate']
    days_needed = (remaining / max_daily_requests) + 1
    
    print(f"📝 Số văn bản còn lại: {remaining}")
    print(f"📊 Limit hàng ngày: {max_daily_requests} requests")
    print(f"📅 Số ngày cần: {days_needed:.1f}")
    print()
    
    # Chia theo ngày
    day = 1
    current_start = 0
    
    while current_start < remaining:
        end = min(current_start + max_daily_requests, remaining)
        count = end - current_start
        
        print(f"📅 Ngày {day}: Dịch {count} văn bản (từ {current_start + 1} đến {end})")
        current_start = end
        day += 1
    
    print(f"\n💡 GỢI Ý:")
    print("1. Chạy script vào mỗi sáng để tận dụng quota mới")
    print("2. Cache sẽ được giữ nguyên, không mất tiến độ")
    print("3. Có thể dừng/tiếp tục bất cứ lúc nào")
    print("4. Consider upgrade lên paid plan để dịch nhanh hơn")

def create_resume_script():
    """Tạo script để tiếp tục dịch thuật ngày mai"""
    resume_script = """#!/usr/bin/env python3
# Script tự động tiếp tục dịch thuật khi quota reset

from japanese_translator_v2 import main

if __name__ == "__main__":
    print("🌅 TIẾP TỤC DỊCH THUẬT - QUOTA MỚI")
    print("=" * 40)
    main()
"""
    
    with open('resume_translation.py', 'w', encoding='utf-8') as f:
        f.write(resume_script)
    
    print("💾 Đã tạo script tiếp tục: resume_translation.py")
    print("   Chạy script này vào ngày mai để tiếp tục!")

def main():
    """Hàm chính để xử lý quota exceeded"""
    try:
        from config import get_config
        config = get_config()
        if not config:
            print("❌ Không tìm thấy config!")
            return
        
        input_file = config['input_file']
        cache_file = config['cache_file']
        
    except ImportError:
        # Fallback config
        input_file = 'kienthucchung.docx'
        cache_file = 'translation_cache_v2.json'
    
    print("🚫 GEMINI API DAILY QUOTA EXCEEDED HANDLER")
    print("=" * 60)
    print("Free tier limit: 50 requests/ngày")
    print("Quota sẽ reset vào 00:00 UTC (07:00 sáng VN)")
    print()
    
    # Phân tích công việc còn lại
    stats = analyze_remaining_work(input_file, cache_file)
    
    if stats:
        # Tạo kế hoạch
        create_daily_plan(stats)
        
        # Tạo script tiếp tục
        create_resume_script()
    
    print("\n🎯 NEXT STEPS:")
    print("1. Đợi đến sáng mai (quota reset 07:00 VN)")
    print("2. Chạy: python resume_translation.py")
    print("3. Hoặc chạy lại: python japanese_translator_v2.py")
    print("4. Cache sẽ được giữ nguyên, tiếp tục từ chỗ dừng")

if __name__ == "__main__":
    main()
