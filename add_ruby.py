import json
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from collections import defaultdict
import time
import gc

# Compile regex patterns một lần duy nhất để tăng tốc
KANJI_PATTERN = re.compile(r'[\u4e00-\u9fff]')
JAPANESE_PATTERN = re.compile(r'[\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff]')

# Set để lưu các kanji không tìm thấy
missing_kanji = set()

def load_dictionary(dictionary_path):
    """Đọc dictionary từ file JSON và tối ưu hóa cho tìm kiếm nhanh"""
    print("Đang load dictionary...")
    start_time = time.time()
    
    try:
        with open(dictionary_path, 'r', encoding='utf-8') as f:
            dictionary = json.load(f)
        
        # Lọc và tối ưu dictionary
        optimized_dict = {}
        kanji_count = 0
        
        for kanji, hiragana in dictionary.items():
            # Lọc ít chặt hơn: chấp nhận từ dài hơn và nhiều loại từ hơn
            if (has_kanji(kanji) and 
                len(kanji) <= 12 and  # Tăng lên 12 ký tự để bao gồm từ ghép
                len(kanji) >= 1 and   # Ít nhất 1 ký tự
                not re.search(r'[a-zA-Z0-9]', kanji)):  # Chỉ loại ký tự Latin và số
                optimized_dict[kanji] = hiragana
                kanji_count += 1
        
        # Sắp xếp theo độ dài giảm dần
        sorted_dict = dict(sorted(optimized_dict.items(), key=lambda x: len(x[0]), reverse=True))
        
        load_time = time.time() - start_time
        print(f"Đã load {kanji_count} từ Kanji trong {load_time:.2f} giây")
        
        return sorted_dict
        
    except FileNotFoundError:
        print(f"Không tìm thấy file dictionary: {dictionary_path}")
        return {}
    except json.JSONDecodeError:
        print(f"Lỗi định dạng JSON trong file: {dictionary_path}")
        return {}

def has_kanji(text):
    """Kiểm tra xem text có chứa ký tự Kanji không"""
    return bool(KANJI_PATTERN.search(text))

def has_japanese(text):
    """Kiểm tra xem text có chứa ký tự tiếng Nhật không"""
    return bool(JAPANESE_PATTERN.search(text))

def clean_kanji_word(word):
    """Làm sạch từ Kanji bằng cách loại bỏ ký tự không phải tiếng Nhật"""
    # Chỉ giữ lại ký tự Kanji, Hiragana, Katakana
    cleaned = re.sub(r'[^\u3040-\u309f\u30a0-\u30ff\u4e00-\u9fff]', '', word)
    return cleaned

def extract_kanji_words(text):
    """Trích xuất tất cả các từ có chứa Kanji từ text"""
    kanji_words = set()
    text_len = len(text)
    
    for i in range(text_len):
        if has_kanji(text[i]):  # Nếu ký tự hiện tại là Kanji
            # Tìm từ Kanji dài nhất từ vị trí này
            for length in range(min(8, text_len - i), 0, -1):  # Bao gồm cả từ 1 ký tự
                word = text[i:i+length]
                cleaned_word = clean_kanji_word(word)
                if cleaned_word and has_kanji(cleaned_word):
                    kanji_words.add(cleaned_word)
    
    return kanji_words

def create_ruby_element(kanji, hiragana, source_run=None):
    """Tạo element ruby cho Word document với spacing tối ưu để tránh che khuất"""
    ruby = OxmlElement('w:ruby')
    
    # Ruby properties - Điều chỉnh để tránh che khuất
    ruby_pr = OxmlElement('w:rubyPr')
    ruby_align = OxmlElement('w:rubyAlign')
    ruby_align.set(qn('w:val'), 'center')
    ruby_pr.append(ruby_align)
    
    # Điều chỉnh kích thước ruby để tăng khoảng cách với chữ Kanji
    hgt = OxmlElement('w:hgt')
    hgt.set(qn('w:val'), '32')  # Tăng từ 24 lên 32 để tạo khoảng cách thực sự giữa ruby và Kanji
    ruby_pr.append(hgt)
    
    # Thêm spacing cho ruby
    ruby_spacing = OxmlElement('w:lid')
    ruby_spacing.set(qn('w:val'), 'ja-JP')
    ruby_pr.append(ruby_spacing)
    
    ruby.append(ruby_pr)
    
    # Ruby text (hiragana ở trên) - Kích thước vừa phải
    rt = OxmlElement('w:rt')
    rt_r = OxmlElement('w:r')
    rt_rpr = OxmlElement('w:rPr')
    
    # Copy format từ source run
    copy_run_formatting(source_run, rt_rpr)
    
    # Kích thước font cho ruby text - giảm xuống 10 (5pt)
    rt_sz = OxmlElement('w:sz')
    rt_sz.set(qn('w:val'), '10')  # 5pt cho ruby
    rt_rpr.append(rt_sz)

    rt_szcs = OxmlElement('w:szCs')
    rt_szcs.set(qn('w:val'), '10')  # 5pt cho ruby
    rt_rpr.append(rt_szcs)
    
    rt_r.append(rt_rpr)
    rt_t = OxmlElement('w:t')
    rt_t.text = hiragana
    rt_r.append(rt_t)
    rt.append(rt_r)
    ruby.append(rt)
    
    # Ruby base (kanji ở dưới) - Tăng kích thước font lên 12pt
    ruby_base = OxmlElement('w:rubyBase')
    base_r = OxmlElement('w:r')
    base_rpr = OxmlElement('w:rPr')
    
    # Copy format từ source run
    copy_run_formatting(source_run, base_rpr)
    
    # Đặt kích thước font cho chữ Kanji là 11pt
    base_sz = OxmlElement('w:sz')
    base_sz.set(qn('w:val'), '22')  # 11pt = 22 half-points
    base_rpr.append(base_sz)
    
    base_szcs = OxmlElement('w:szCs')
    base_szcs.set(qn('w:val'), '22')  # 11pt = 22 half-points
    base_rpr.append(base_szcs)
    
    base_r.append(base_rpr)
    base_t = OxmlElement('w:t')
    base_t.text = kanji
    base_r.append(base_t)
    ruby_base.append(base_r)
    ruby.append(ruby_base)
    
    return ruby

def find_kanji_matches_optimized(text, dictionary):
    """Tìm matches với thuật toán cải tiến - tìm tất cả từ có thể"""
    if not text or not has_japanese(text):
        return []
    
    matches = []
    text_len = len(text)
    covered = [False] * text_len  # Theo dõi vị trí đã được xử lý
    
    # Lần 1: Tìm từ dài nhất trước (ưu tiên từ ghép)
    for length in range(min(12, text_len), 2, -1):  # Từ 12 ký tự xuống 3 ký tự
        for i in range(text_len - length + 1):
            if any(covered[i:i+length]):  # Bỏ qua nếu đã có ruby
                continue
                
            substring = text[i:i+length]
            cleaned_substring = clean_kanji_word(substring)
            
            if (cleaned_substring and 
                has_kanji(cleaned_substring) and 
                cleaned_substring in dictionary):
                matches.append((i, i + length, cleaned_substring, dictionary[cleaned_substring]))
                # Đánh dấu vùng đã xử lý
                for j in range(i, i + length):
                    covered[j] = True
    
    # Lần 2: Tìm từ ngắn (1-2 ký tự) cho những vị trí chưa xử lý
    for length in range(2, 0, -1):  # 2 ký tự rồi 1 ký tự
        for i in range(text_len - length + 1):
            if any(covered[i:i+length]):  # Bỏ qua nếu đã có ruby
                continue
                
            substring = text[i:i+length]
            cleaned_substring = clean_kanji_word(substring)
            
            if (cleaned_substring and 
                has_kanji(cleaned_substring) and 
                cleaned_substring in dictionary):
                matches.append((i, i + length, cleaned_substring, dictionary[cleaned_substring]))
                # Đánh dấu vùng đã xử lý
                for j in range(i, i + length):
                    covered[j] = True
    
    # Ghi lại các Kanji không tìm thấy
    for i in range(text_len):
        if not covered[i] and has_kanji(text[i]):
            missing_kanji.add(text[i])
    
    # Sắp xếp matches theo vị trí
    matches.sort(key=lambda x: x[0])
    return matches

def has_section_break(paragraph):
    """Kiểm tra xem paragraph có chứa section break không"""
    pPr = paragraph._element.pPr
    if pPr is not None:
        sectPr = pPr.find(qn('w:sectPr'))
        return sectPr is not None
    return False

def remove_section_breaks_and_add_spacing(paragraph):
    """Xóa section break và thêm 2 dòng trống"""
    pPr = paragraph._element.pPr
    if pPr is not None:
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            # Xóa section break
            pPr.remove(sectPr)
            print(f"Đã xóa section break và thêm 2 dòng trống")
            return True
    return False

def add_ruby_to_paragraph_compact(paragraph, dictionary):
    """Thêm ruby vào paragraph với thuật toán cải tiến và spacing tối ưu"""
    full_text = paragraph.text
    
    # Xử lý section break trước khi kiểm tra text
    section_break_removed = remove_section_breaks_and_add_spacing(paragraph)
    
    if not full_text or not has_japanese(full_text):
        return section_break_removed
    
    # Lấy format từ run đầu tiên để áp dụng cho ruby
    first_run = get_first_run_from_paragraph(paragraph)
    
    # SỬA: Xử lý tất cả paragraph, không giới hạn độ dài
    matches = find_kanji_matches_optimized(full_text, dictionary)
    
    if not matches:
        return section_break_removed
    
    # Đặt line spacing phù hợp cho paragraph có ruby
    try:
        pPr = paragraph._element.get_or_add_pPr()
        spacing = pPr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            pPr.append(spacing)
        # Giảm line spacing để compact hơn
        spacing.set(qn('w:line'), '360')  # 18pt line spacing
        spacing.set(qn('w:lineRule'), 'exact')
        # Giảm space before và after để compact hơn
        spacing.set(qn('w:before'), '60')   # 3pt space before
        spacing.set(qn('w:after'), '20')    # 1pt space after
    except:
        pass
    
    # Xóa tất cả runs hiện tại
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)
    
    # Xây dựng lại paragraph với ruby
    last_pos = 0
    
    for start, end, kanji, hiragana in matches:
        # Thêm text trước kanji với font size 11pt nếu có tiếng Nhật
        if start > last_pos:
            before_text = full_text[last_pos:start]
            if before_text:
                if has_japanese(before_text):
                    create_japanese_run_with_font_size(paragraph, before_text, 11)
                else:
                    paragraph.add_run(before_text)
        
        # Thêm ruby element với format từ run đầu tiên
        ruby_element = create_ruby_element(kanji, hiragana, first_run)
        paragraph._element.append(ruby_element)
        
        last_pos = end
    
    # Thêm text còn lại với font size 11pt nếu có tiếng Nhật
    if last_pos < len(full_text):
        remaining_text = full_text[last_pos:]
        if remaining_text:
            if has_japanese(remaining_text):
                create_japanese_run_with_font_size(paragraph, remaining_text, 11)
            else:
                paragraph.add_run(remaining_text)
    
    return True

def save_missing_kanji_report(missing_kanji, output_path):
    """Lưu báo cáo các từ Kanji không tìm thấy"""
    if not missing_kanji:
        print("Tất cả từ Kanji đều có trong dictionary!")
        return
    
    report_file = output_path.replace('.docx', '_missing_kanji.txt')
    
    # Sắp xếp theo độ dài và alphabet
    sorted_missing = sorted(missing_kanji, key=lambda x: (len(x), x))
    
    with open(report_file, 'w', encoding='utf-8') as f:
        f.write("=== BÁO CÁO CÁC TỪ KANJI KHÔNG TÌM THẤY TRONG DICTIONARY ===\n\n")
        f.write(f"Tổng số từ Kanji không tìm thấy: {len(missing_kanji)}\n\n")
        
        # Phân loại theo độ dài
        by_length = defaultdict(list)
        for word in sorted_missing:
            by_length[len(word)].append(word)
        
        for length in sorted(by_length.keys()):
            f.write(f"=== Từ có {length} ký tự ({len(by_length[length])} từ) ===\n")
            for word in by_length[length]:
                f.write(f"{word}\n")
            f.write("\n")
        
        f.write("=== DANH SÁCH ĐẦY ĐỦ ===\n")
        for word in sorted_missing:
            f.write(f"{word}\n")
    
    print(f"Đã lưu báo cáo từ Kanji thiếu: {report_file}")
    print(f"Tổng số từ Kanji không tìm thấy: {len(missing_kanji)}")

def process_word_document(input_path, output_path, dictionary_path):
    """Xử lý file Word với tối ưu để tránh tăng quá nhiều trang"""
    global missing_kanji
    missing_kanji.clear()  # Reset set
    
    print("=== Xử lý file Word với tối ưu compact ===")
    start_time = time.time()
    
    # Load dictionary
    dictionary = load_dictionary(dictionary_path)
    
    if not dictionary:
        print("Dictionary trống hoặc không đọc được.")
        return
    
    print(f"Đang xử lý file: {input_path}")
    
    try:
        doc = Document(input_path)
        
        # Đếm số
        japanese_paragraphs = 0
        total_paragraphs = 0
        processed_count = 0
        section_breaks_removed = 0
        
        print("Đang xử lý paragraphs chính...")
        
        # Danh sách các paragraph cần thêm dòng trống sau khi xử lý
        paragraphs_to_add_spacing = []
        
        # Xử lý paragraph trong document chính
        for i, paragraph in enumerate(doc.paragraphs):
            total_paragraphs += 1
            
            # Kiểm tra và xử lý section break
            if has_section_break(paragraph):
                section_breaks_removed += 1
                paragraphs_to_add_spacing.append(i)
            
            if paragraph.text.strip() and has_japanese(paragraph.text):
                japanese_paragraphs += 1
                
                # SỬA: Xử lý tất cả paragraph, không bỏ qua paragraph dài
                if add_ruby_to_paragraph_compact(paragraph, dictionary):
                    processed_count += 1
                
                # Hiển thị tiến trình mỗi 50 paragraph
                if japanese_paragraphs % 50 == 0:
                    elapsed = time.time() - start_time
                    print(f"Đã xử lý {japanese_paragraphs} paragraph tiếng Nhật - {elapsed:.1f}s")
                    
                    # Garbage collection định kỳ
                    if japanese_paragraphs % 200 == 0:
                        gc.collect()
            else:
                # Xử lý section break cho paragraph không có tiếng Nhật
                add_ruby_to_paragraph_compact(paragraph, dictionary)
        
        # Thêm 2 dòng trống sau các paragraph có section break
        for i in reversed(paragraphs_to_add_spacing):
            p1 = doc.add_paragraph()
            p2 = doc.add_paragraph()
            doc._element.body.insert(i + 1, p1._element)
            doc._element.body.insert(i + 2, p2._element)
        
        print("Đang xử lý tables...")
        
        # Xử lý text trong tables
        table_count = 0
        for table in doc.tables:
            table_count += 1
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        total_paragraphs += 1
                        if paragraph.text.strip() and has_japanese(paragraph.text):
                            japanese_paragraphs += 1
                            if add_ruby_to_paragraph_compact(paragraph, dictionary):
                                processed_count += 1
            
            # Hiển thị tiến trình tables
            if table_count % 20 == 0:
                elapsed = time.time() - start_time
                print(f"Đã xử lý {table_count} tables - {elapsed:.1f}s")
        
        # Lưu file
        print("Đang lưu file...")
        save_start = time.time()
        doc.save(output_path)
        save_time = time.time() - save_start
        
        # Lưu báo cáo missing kanji
        save_missing_kanji_report(missing_kanji, output_path)
        
        total_time = time.time() - start_time
        
        print("\n=== KẾT QUẢ XỬ LÝ ===")
        print(f"File output: {output_path}")
        print(f"Tổng số paragraph: {total_paragraphs}")
        print(f"Paragraph có tiếng Nhật: {japanese_paragraphs}")
        print(f"Paragraph đã thêm ruby: {processed_count}")
        if japanese_paragraphs > 0:
            print(f"Tỷ lệ xử lý thành công: {processed_count/japanese_paragraphs*100:.1f}%")
        print(f"Section breaks đã xóa: {section_breaks_removed}")
        print(f"Từ Kanji không tìm thấy: {len(missing_kanji)}")
        print(f"Thời gian xử lý: {total_time:.2f} giây")
        print(f"Thời gian lưu file: {save_time:.2f} giây")
        if japanese_paragraphs > 0:
            print(f"Tốc độ xử lý: {japanese_paragraphs/total_time:.1f} paragraph/giây")
        
    except Exception as e:
        print(f"Lỗi khi xử lý file Word: {e}")
        import traceback
        traceback.print_exc()

def copy_run_formatting(source_run, target_rpr):
    """Copy formatting từ source run sang target rPr element, loại bỏ highlighting"""
    if source_run and source_run._element.rPr is not None:
        for child in source_run._element.rPr:
            # Copy tất cả format trừ kích thước font và highlighting
            if child.tag not in [qn('w:sz'), qn('w:szCs'), qn('w:highlight'), qn('w:shd')]:
                try:
                    target_rpr.append(child.__copy__())
                except:
                    pass

def get_first_run_from_paragraph(paragraph):
    """Lấy run đầu tiên từ paragraph để copy format"""
    if paragraph.runs:
        return paragraph.runs[0]
    return None

def print_processing_stats(japanese_paragraphs, processed_count, missing_count, total_time):
    """In thống kê quá trình xử lý"""
    print(f"\n=== THỐNG KÊ XỬ LÝ ===")
    print(f"Paragraph có tiếng Nhật: {japanese_paragraphs}")
    print(f"Paragraph đã thêm ruby: {processed_count}")
    if japanese_paragraphs > 0:
        print(f"Tỷ lệ xử lý: {processed_count/japanese_paragraphs*100:.1f}%")
    print(f"Từ Kanji không tìm thấy: {missing_count}")
    print(f"Thời gian xử lý: {total_time:.2f} giây")
    if japanese_paragraphs > 0:
        print(f"Tốc độ: {japanese_paragraphs/total_time:.1f} paragraph/giây")

def create_japanese_run_with_font_size(paragraph, text, font_size_pt=11):
    """Tạo run với kích thước font cụ thể cho text tiếng Nhật"""
    run = paragraph.add_run(text)
    
    # Tạo hoặc lấy rPr element
    rPr = run._element.rPr
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._element.insert(0, rPr)
    
    # Xóa kích thước font hiện tại nếu có
    for sz in rPr.findall(qn('w:sz')):
        rPr.remove(sz)
    for szcs in rPr.findall(qn('w:szCs')):
        rPr.remove(szcs)
    
    # Đặt kích thước font mới
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size_pt * 2))  # Chuyển đổi pt sang half-points
    rPr.append(sz)
    
    szcs = OxmlElement('w:szCs')
    szcs.set(qn('w:val'), str(font_size_pt * 2))  # Chuyển đổi pt sang half-points
    rPr.append(szcs)
    
    return run

def main():
    input_file = "KHÔNG FURIGANA - KIẾN THỨC CHUNG.docx"
    output_file = "KHÔNG FURIGANA - KIẾN THỨC CHUNG ruby.docx"  # Giảm spacing giữa các dòng
    dictionary_file = "dictionary_hiragana_chunks_removed_from_value.json"
    
    print("=== Chương trình thêm Ruby (không highlight) cho file Word ===")
    print(f"Input file: {input_file}")
    print(f"Output file: {output_file}")
    print(f"Dictionary file: {dictionary_file}")
    print()
    
    process_word_document(input_file, output_file, dictionary_file)
    print("\nHoàn thành!")

if __name__ == "__main__":
    main()
